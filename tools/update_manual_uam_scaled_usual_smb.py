from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 SMB time windows.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 UAM scaled usual SMB.docx")

doc = Document(SOURCE)

output_paragraph = next(
    p for p in doc.paragraphs if p.text.startswith("For output, the calculation compares")
)
output_paragraph.text = (
    "For output, Tier 3 now uses the already-calculated usual current SMB as an explicit "
    "assessment input. That usual SMB already includes the live variable SMB-delivery ratio "
    "and any active BolusGiven or BolusGivenMild ratio increase. Tier 3 multiplies it by the "
    "effective Tier 3 boost scale and compares that result with both its basal-based scaled "
    "candidate and its baseline-ratio candidate. In compact form, the provisional Tier 3 "
    "amount is the largest of: scaled current basal, baseline-ratio SMB, and usual current "
    "SMB × Tier 3 scale. This assessment runs only with the Virtual Pump. Tier 3 is considered "
    "to have enhanced the candidate only when the rounded result exceeds the ordinary pre-boost "
    "SMB. The result remains limited by the configured maximum boost bolus and the remaining "
    "percentage-based IOB allowance; that IOB allowance is enforced again immediately before "
    "final pump-increment rounding."
)

relationship = next(
    p for p in doc.paragraphs if p.text.startswith("Relationship to SMBs already delivered.")
)
relationship.text = (
    "Relationship to SMBs already delivered. Tier 3's initial maximum boost bolus is an "
    "absolute per-SMB cap, not a percentage of recently delivered SMBs. Its new scaled-usual-SMB "
    "branch increases the current ordinary SMB by the effective Tier 3 scale; it does not discard "
    "an increase already produced by the live SMB-delivery ratio or by BolusGiven/Mild. For example, "
    "if the usual current SMB is 0.40 U and the effective Tier 3 scale is 1.5, that branch proposes "
    "0.60 U before the Tier 3 ceilings and shared safety reductions. Delivered SMBs affect current "
    "IOB once their bolus records arrive. After Tier 3 proposes an amount, the shared safety path "
    "directly uses delivered-SMB history: rolling 10-minute limits are 0.6 U from 00:30–04:00 and "
    "1.5 U otherwise; the rolling 30-minute ceiling is 2.1 U; late fast-rise delivery is reduced "
    "to 75% after 1.5 U/30 min and 50% after 1.9 U/30 min; rapid spacing and the sub-7.5 mmol/L "
    "heavy-delivery cooldown may further reduce or zero the candidate."
)

doc.save(OUTPUT)
print(OUTPUT)
