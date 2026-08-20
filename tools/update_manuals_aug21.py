from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PY_OUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320")
OPS_SOURCE = Path(r"C:\winword\aaa\AutoISF Idiosyncrasies and Operations Manual mydoc Aug 21 26 Tier3 added.docx")
OPS_INTERMEDIATE = PY_OUT / "AutoISF_mydoc_boosts_intermediate.docx"
SETTINGS_SOURCE = Path(r"C:\winword\aaa\AutoISF manual -settings-reference claude Monday, August 17th, 2026 mydoc_Aug 19 26_DELL17.docx")
SETTINGS_OUTPUT = PY_OUT / "AutoISF_Settings_mydoc_current_column_intermediate.docx"


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def strip_generated_navigation(doc):
    body = doc._element.body
    title = next(p for p in doc.paragraphs if p.style.name == "Title")
    for child in list(body):
        if child is title._element:
            break
        remove_element(child)
    for hyperlink in list(body.iter(qn("w:hyperlink"))):
        visible = "".join(t.text or "" for t in hyperlink.iter(qn("w:t")))
        if visible.strip() == "Back to TOC":
            remove_element(hyperlink)
    for bookmark in list(body.iter(qn("w:bookmarkStart"))):
        remove_element(bookmark)
    for bookmark in list(body.iter(qn("w:bookmarkEnd"))):
        remove_element(bookmark)
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            style = paragraph.style.name
            paragraph.text = paragraph.text.strip()
            paragraph.style = doc.styles[style]


def insert_labeled_paragraph(doc, anchor, label, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.add_run(label).bold = True
    paragraph.add_run(text)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    anchor._element.addprevious(paragraph._element)


# ---------- Operations Manual: expand Tier 3 relative to the other boosts ----------
ops = Document(OPS_SOURCE)
strip_generated_navigation(ops)
anchor = next(p for p in ops.paragraphs if p.style.name == "Heading 2" and p.text.startswith("c. The weight system"))

insert_labeled_paragraph(
    ops,
    anchor,
    "Relationship to SMBs already delivered. ",
    "Tier 3's initial maximum boost bolus is an absolute per-SMB cap, not a percentage of recently delivered SMBs. Its ordinary candidate is directly affected by the current smb_delivery_ratio, and delivered SMBs affect current IOB once their bolus records arrive. After Tier 3 proposes an amount, the shared safety path directly uses delivered-SMB history: rolling 10-minute limits are 0.6 U from 00:30–04:00 and 1.5 U otherwise; the rolling 30-minute ceiling is 2.1 U; late fast-rise delivery is reduced to 75% after 1.5 U/30 min and 50% after 1.9 U/30 min; rapid spacing and the sub-7.5 mmol/L heavy-delivery cooldown may further reduce or zero the candidate."
)
insert_labeled_paragraph(
    ops,
    anchor,
    "BolusGiven strong boost. ",
    "This is separate from Tier 3 and is controlled by Enable BolusGiven/BolusGivenMild automations. Branches 1 and 2 react after a normal bolus, using bolus age, COB, glucose rise, raw-Libre confirmation, activity, profile and state gates. Branch 3 is delivery-driven: normally no bolus or carbs for at least 120 minutes, a sufficiently large five-minute IOB rise, strong AAPS/raw glucose rise, BG no higher than 9.5 mmol/L, low activity, not on the Low/MJ profile, and not literally MJ active. Branch 3 also has a raw-signal bypass when recent SMB delivery was suppressed despite a strong continuing rise."
)
insert_labeled_paragraph(
    ops,
    anchor,
    "Strong-boost effects. ",
    "When a BolusGiven branch fires it cancels the current temporary target, sets iob_threshold_percent to 71, selects the Standard Profile, restores acceleration weight to its configured normal value, raises postprandial ISF weight, sets SMB delivery ratio to Mild Boost base + 0.03, applies a 110% profile for two minutes, and applies a 4.2 mmol/L target for two minutes. Branch 3 also arms a 30-minute recent-boost marker that can restore selected fast-rise-capped SMBs to the saved uncapped amount; recent-low and final cumulative-delivery protections run afterward and still take precedence."
)
insert_labeled_paragraph(
    ops,
    anchor,
    "BolusGivenMild and Mild Failsafe. ",
    "Mild covers a lower raw-delta band than strong branch 3 and is designed to be mutually exclusive with it. It uses delivery change, raw/AAPS rise confirmation, no recent bolus/carbs, low activity, profile/MJ gates and cross-cooldowns. It leaves profile percentage and IOB threshold unchanged, but raises smb_delivery_ratio to Mild base +0.05 below 7.5 mmol/L, +0.02 below 9.0 mmol/L, or the base at/above 9.0 mmol/L; it also sets a 5.0 mmol/L target for two minutes and raises postprandial ISF weight. Mild Failsafe applies similar effects when BG is above 6.5 mmol/L, the rise is confirmed across delta horizons, IOB is at most 0.20 U, and no SMB has been delivered for 20 minutes."
)
insert_labeled_paragraph(
    ops,
    anchor,
    "Initiation order relative to Tier 3. ",
    "BolusGiven/Mild automations run first and can change smb_delivery_ratio, profile percentage, target and weights. AutoISF then builds the profile passed into DetermineBasal, where Tier 3 independently evaluates its own Virtual-Pump, delta-ratio, BG, eventual-BG, positive-insulin-requirement and IOB-ceiling criteria. A prior boost can therefore alter Tier 3's ordinary candidate, effective scale or target context, but it does not automatically start Tier 3. Tier 3's separate switch and all of its own gates must still pass; afterward every common SMB safety modifier and the normal SMB interval still apply."
)
ops.save(OPS_INTERMEDIATE)


# ---------- Settings Reference: merge Type and Current tables ----------
settings = Document(SETTINGS_SOURCE)
current_table = settings.tables[5]
current_values = {}
for row in current_table.rows[1:]:
    name = row.cells[0].text.strip()
    current_values[name] = row.cells[2].text.strip()

for table in settings.tables[:5]:
    table.cell(0, 0).text = "Setting"
    table.cell(0, 1).text = "Default (range) [type]"
    table.cell(0, 2).text = "Detailed description"
    table.cell(0, 3).text = "Current setting (Client)"
    for row in table.rows[1:]:
        setting_name = row.cells[0].text.strip()
        setting_type = row.cells[1].text.strip()
        default_range = row.cells[2].text.strip()
        description = row.cells[3].text.strip()
        current = current_values.get(setting_name, "Not in snapshot")
        row.cells[1].text = f"{default_range} [{setting_type}]"
        row.cells[2].text = description
        row.cells[3].text = current

# Remove the now-redundant standalone Current Settings section and comparison table.
start = next(p for p in settings.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Current Settings"))
end = next(p for p in settings.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Not Yet In This Snapshot"))
removing = False
for paragraph in list(settings.paragraphs):
    if paragraph._element is start._element:
        removing = True
    if paragraph._element is end._element:
        removing = False
    if removing:
        remove_element(paragraph._element)
remove_element(current_table._element)

# Update the pre-snapshot Tier 3 IOB wording to the implemented percentage preference.
for paragraph in settings.paragraphs:
    if paragraph.text.startswith("• UAM Boost max IOB"):
        paragraph.text = "• UAM Boost max IOB (% of max_iob) — default 10% (1–100%), no live value available from this build's export."

# Apply explicit four-column geometry to the main tables.
for table in settings.tables[:5]:
    section = settings.sections[0]
    usable_twips = int((section.page_width - section.left_margin - section.right_margin) / 635)
    widths = [int(usable_twips * p) for p in (0.22, 0.20, 0.43)]
    widths.append(usable_twips - sum(widths))
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(usable_twips))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))

settings.save(SETTINGS_OUTPUT)
print(OPS_INTERMEDIATE)
print(SETTINGS_OUTPUT)
