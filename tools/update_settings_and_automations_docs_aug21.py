import ast
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320")
SETTINGS_SOURCE = Path(r"C:\winword\aaa\AutoISF Settings Reference mydoc Aug 21 26 current column.docx")
SETTINGS_OUTPUT = ROOT / "AutoISF Settings Reference mydoc Aug 21 26 List1 List2 commands added.docx"
AUTOMATIONS_SOURCE = Path(r"C:\winword\aaa\AutoISF_Automations_List5 mydoc.docx")
AUTOMATIONS_OUTPUT = ROOT / "AutoISF Automations List mydoc Aug 21 26 current code registry.docx"


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths, font_size=8.2):
    table.autofit = False
    table.style = "Normal Table"
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "A6A6A6")
        borders.append(border)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    header_pr.append(header)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(1.5)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def load_manual_command_rows():
    source = (ROOT / "tools" / "add_list_tt_catalogue_main_manual_aug21.py").read_text(encoding="utf-8")
    tree = ast.parse(source)
    found = {}
    for node in tree.body:
        if isinstance(node, ast.Assign) and len(node.targets) == 1 and isinstance(node.targets[0], ast.Name):
            if node.targets[0].id in {"list1_rows", "list2_rows"}:
                found[node.targets[0].id] = ast.literal_eval(node.value)
    return found["list1_rows"], found["list2_rows"]


def append_settings_catalogue():
    list1_rows, list2_rows = load_manual_command_rows()
    doc = Document(SETTINGS_SOURCE)
    # Add explicit sections at the end; this preserves all existing tables and the Current setting column.
    heading = doc.add_paragraph("List 1 and List 2 command/temporary-target reference", style="Heading 1")
    for run in heading.runs:
        run.font.underline = True
    doc.add_paragraph(
        "These mmol/L values are command signals, not clinical targets. Pump/Virtual builds execute the selected handler locally without creating a temporary target. AAPSClient uses a five-minute relay TT so the pump can receive the command; the pump applies the action and cancels the relay. List 2 graph rows are local display controls and never create relay targets."
    )

    h1 = doc.add_paragraph("a. List 1 - settings and relay commands", style="Heading 2")
    h1.paragraph_format.left_indent = Pt(18)
    table1 = doc.add_table(rows=1, cols=4)
    for cell, value in zip(table1.rows[0].cells, ["Command", "Relay TT", "Effect", "Current-setting behavior"]):
        cell.text = value
    for command, code, effect in list1_rows:
        row = table1.add_row().cells
        row[0].text = command
        row[1].text = code
        row[2].text = effect
        row[3].text = "Pump/Virtual: local value shown and changed after confirmation. Client: mirrored pump value shown when available; local Client preference is not changed."
    configure_table(table1, [1850, 1100, 3760, 2650], 7.7)

    h2 = doc.add_paragraph("b. List 2 - actions, relays and local graph controls", style="Heading 2")
    h2.paragraph_format.left_indent = Pt(18)
    table2 = doc.add_table(rows=1, cols=4)
    for cell, value in zip(table2.rows[0].cells, ["Command", "Relay TT", "Effect", "Transport/status"]):
        cell.text = value
    for command, code, effect in list2_rows:
        row = table2.add_row().cells
        row[0].text = command
        row[1].text = code
        row[2].text = effect
        row[3].text = "Graph rows: local only. Clinical actions: direct on Pump/Virtual; five-minute relay on Client after confirmation. AnyDesk preserves an existing real TT by using Note-only transport."
    configure_table(table2, [1850, 1100, 3760, 2650], 7.7)
    doc.save(SETTINGS_OUTPUT)


def parse_registry_keys():
    text = (ROOT / "core" / "utils" / "src" / "main" / "kotlin" / "app" / "aaps" / "core" / "utils" / "CodedAutomationNames.kt").read_text(encoding="utf-8")
    body = re.search(r"val KEYS: List<String> = listOf\((.*?)\n\s*\)", text, re.S).group(1)
    return re.findall(r'"([^"]+)"', body)


def automation_category(key):
    if key.endswith("TT") or key.startswith("TodOffset") or "Weight" in key or key.startswith("WizardPct") or key.startswith("Smb") or key.startswith("AutoIsfMax") or key.startswith("PeakInsulin"):
        return "Command / setting relay"
    if key in {"MJ4", "MJ5", "Test3"}:
        return "Validation/test"
    if key in {"AlarmHypo1", "AlarmHypo2", "GentleHypoRisk", "PrepareSet50", "SkittlesHypoRisk", "Extra50", "iobTHDaytimeFloor"}:
        return "Hypoglycaemia protection"
    if key.startswith("Battery") or "Pod" in key or key.startswith("Sensor") or key.startswith("PreSoak") or key == "ConnectPod":
        return "Device/pump/sensor"
    if key.startswith("Bolus") or key == "VirtualPseudoWizard":
        return "Bolus/post-bolus"
    if key.startswith("MJ") or key == "MoreMJ":
        return "Mounjaro state"
    if "Steroid" in key:
        return "Steroid state"
    return "Operational automation"


def extract_code_summary(key, lines):
    special = {
        "LibreUkf2ToggleTT": "Retired/unassigned legacy relay key. Code 5.154 was freed when UKF2 graph history was made continuously current without a graph checkbox changing the live dosing engine. No active List/TT handler presently uses it; it remains in the coded-name registry so an old native automation title can still be classified for review.",
    }
    if key in special:
        return special[key]
    exact_ready = re.compile(rf'readyToRun\("{re.escape(key)}"\s*[,)]')
    exact_mark = re.compile(rf'markRun\("{re.escape(key)}"\s*\)')
    patterns = [exact_ready, exact_mark]
    occurrence = None
    for pattern in patterns:
        candidates = [i for i, line in enumerate(lines) if pattern.search(line)]
        occurrence = next((i for i in candidates if "if (" in lines[i] or lines[i].lstrip().startswith("if")), None)
        if occurrence is None:
            occurrence = candidates[0] if candidates else None
        if occurrence is not None:
            break
    if occurrence is None:
        exact_literal = re.compile(rf'"{re.escape(key)}"')
        occurrence = next((i for i, line in enumerate(lines) if exact_literal.search(line)), None)
    if occurrence is None:
        return "Registry key is present, but no direct handler occurrence was found by the documentation extractor; inspect the current source before use."
    window_start = max(0, occurrence - 90)
    banner = None
    for i in range(window_start, occurrence):
        stripped = lines[i].strip()
        if stripped.startswith("// ---") or stripped.startswith("// Code port") or stripped.startswith("// Standalone"):
            banner = i
    start = banner if banner is not None else max(window_start, occurrence - 18)
    comments = []
    for line in lines[start:occurrence]:
        stripped = line.strip()
        if stripped.startswith("//"):
            value = stripped[2:].strip().strip("-").strip()
            if value and not value.startswith(("TODO", "===", "***")):
                comments.append(value)
    summary = " ".join(comments)
    summary = re.sub(r"\s+", " ", summary).strip()
    if not summary:
        if key.startswith("TodOffset"):
            summary = "Member of the eight TodOffset TT pairs. A five-minute relay guard recognizes its exact 5.092-5.136 mmol command, changes the corresponding fixed time-of-day varOffset by 0.1 within the +/-2.0 clamp, cancels the command TT and records/notifies the result."
        else:
            summary = "Current handler is defined in OpenAPSAutoISFPlugin.kt; the adjacent source block contains no standalone prose comment."
    # Keep enough criteria/action detail while avoiding multi-page cells.
    if len(summary) > 900:
        cut = summary.rfind(". ", 0, 900)
        summary = summary[: cut + 1 if cut > 300 else 900].rstrip() + ("" if cut > 300 else "...")
    return summary


def rebuild_automations_list():
    keys = parse_registry_keys()
    source_path = ROOT / "plugins" / "aps" / "src" / "main" / "kotlin" / "app" / "aaps" / "plugins" / "aps" / "openAPSAutoISF" / "OpenAPSAutoISFPlugin.kt"
    lines = source_path.read_text(encoding="utf-8").splitlines()
    doc = Document(AUTOMATIONS_SOURCE)
    # Remove the obsolete table and all old body paragraphs, preserving the document's section/layout container.
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    title = doc.add_paragraph("AutoISF Coded Automations - Current Code Registry", style="Heading 1")
    for run in title.runs:
        run.font.underline = True
    doc.add_paragraph(
        f"Current-code audit dated 2026-08-21. This document lists all {len(keys)} keys in CodedAutomationNames.KEYS, in registry order. The summary column is extracted from the comments attached to each current readyToRun()/markRun() handler in OpenAPSAutoISFPlugin.kt; it replaces the obsolete 56-row snapshot."
    )
    doc.add_paragraph(
        "Registry meaning. EXACT and CLOSE matching are used only to decide whether a native Automation-tab item plausibly duplicates a coded automation. Non-matching native automations are not suppressed. Command/setting-relay rows are included because they are coded automation handlers even though their trigger is a List/TT command rather than an autonomous glucose rule. Disabled handlers remain listed and are labelled by their current source comments."
    )
    table = doc.add_table(rows=1, cols=4)
    for cell, value in zip(table.rows[0].cells, ["#", "Current coded key", "Type", "Current code definition / criteria and action"]):
        cell.text = value
    for index, key in enumerate(keys, 1):
        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = key
        row[2].text = automation_category(key)
        row[3].text = extract_code_summary(key, lines)
    configure_table(table, [500, 2050, 1600, 5210], 7.4)
    doc.save(AUTOMATIONS_OUTPUT)
    return len(keys)


append_settings_catalogue()
count = rebuild_automations_list()
print(SETTINGS_OUTPUT)
print(AUTOMATIONS_OUTPUT)
print(f"automation_keys={count}")
