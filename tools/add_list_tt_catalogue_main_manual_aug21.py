from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 UKF and battery added.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 UKF battery and TT commands.docx")
doc = Document(SOURCE)


def next_bookmark_id():
    values = []
    for node in doc.element.body.xpath(".//w:bookmarkStart"):
        try:
            values.append(int(node.get(qn("w:id"))))
        except (TypeError, ValueError):
            pass
    return max(values, default=0) + 1


bookmark_id = next_bookmark_id()


def add_bookmark(paragraph, name):
    global bookmark_id
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(1, start)
    paragraph._p.append(end)
    bookmark_id += 1


def add_internal_link(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([props, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_paragraph_before(reference, text="", style=None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    reference._element.addprevious(p._element)
    return p


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "A6A6A6")
        borders.append(border)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    header_pr.append(header)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True


def add_command_table(reference, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Normal Table"
    table.rows[0].cells[0].text = "Command"
    table.rows[0].cells[1].text = "Relay TT (mmol/L)"
    table.rows[0].cells[2].text = "Current effect"
    for command, code, effect in rows:
        cells = table.add_row().cells
        cells[0].text = command
        cells[1].text = code
        cells[2].text = effect
    configure_table(table, [2300, 1500, 5560])
    reference._element.addprevious(table._element)
    return table


list1_rows = [
    ("SMBdel base + mild-Bst", "5.002 / 5.004", "Decrease/increase both SMB-delivery baseline and Mild Boost ratio by 0.01."),
    ("Toggle Libre sensor adjustment", "5.006", "Toggle aging-sensor/Libre adjustment on or off."),
    ("Toggle boost automations", "5.008", "Toggle the whole boost-automation group on or off."),
    ("pp ISF weight - normal", "5.012 / 5.014", "Decrease/increase normal postprandial ISF weight by 0.01."),
    ("Acceleration ISF weight - normal", "5.016 / 5.018", "Decrease/increase normal acceleration ISF weight by 0.05."),
    ("Duration ISF weight - normal", "5.022 / 5.024", "Decrease/increase normal duration ISF weight by 0.1."),
    ("Libre slope - original", "5.026 / 5.028", "Decrease/increase the baseline Libre calibration slope by 0.01."),
    ("Libre offset - original", "5.032 / 5.034", "Decrease/increase the baseline Libre calibration offset by 0.05."),
    ("SMB offset", "5.036 / 5.038", "Decrease/increase the hard SMB-offset override by 0.1."),
    ("Clean graph view", "5.042", "Hide SMB dose labels/arrows and show the plain solid-green view."),
    ("Wizard bolus percentage", "5.046 / 5.048", "Decrease/increase the Wizard percentage by five percentage points."),
    ("Mild Boost ratio", "5.052 / 5.054", "Decrease/increase Mild Boost ratio alone by 0.01."),
    ("pp ISF weight - high", "5.056 / 5.058", "Decrease/increase boosted postprandial ISF weight by 0.01."),
    ("Acceleration ISF weight - high", "5.062 / 5.064", "Decrease/increase boosted acceleration ISF weight by 0.01."),
    ("Higher-ISF-range weight", "5.068 / 5.070", "Decrease/increase high-BG ISF weight by 0.1."),
    ("Peak insulin time", "5.074 / 5.076", "Decrease/increase insulin peak by five minutes."),
    ("AutoISF maximum - low BG", "5.080 / 5.082", "Decrease/increase the low-BG AutoISF maximum by 0.1."),
    ("AutoISF maximum - normal", "5.086 / 5.088", "Decrease/increase the normal AutoISF maximum by 0.1."),
    ("TOD offset 00:00-02:00", "5.092 / 5.094", "Decrease/increase this time-of-day offset by 0.1."),
    ("TOD offset 02:00-04:00", "5.098 / 5.100", "Decrease/increase this time-of-day offset by 0.1."),
    ("TOD offset 04:00-06:00", "5.104 / 5.106", "Decrease/increase this time-of-day offset by 0.1."),
    ("TOD offset 06:00-09:00", "5.110 / 5.112", "Decrease/increase this time-of-day offset by 0.1."),
    ("TOD offset 09:00-12:00", "5.116 / 5.118", "Decrease/increase this time-of-day offset by 0.1."),
    ("TOD offset 12:00-18:00", "5.122 / 5.124", "Decrease/increase this time-of-day offset by 0.1."),
    ("TOD offset 18:00-22:00", "5.128 / 5.130", "Decrease/increase this time-of-day offset by 0.1."),
    ("TOD offset 22:00-00:00", "5.134 / 5.136", "Decrease/increase this time-of-day offset by 0.1."),
    ("Toggle Graph2 carb-model curve", "5.138", "Toggle the empirical/theoretical carb-model curve display."),
    ("Cloud logs upload", "5.140", "Zip and send logs to configured cloud storage, otherwise use email."),
    ("Toggle Graph5", "5.142", "Toggle the Graph5 comparison panel."),
    ("MJ state manual override", "5.144 / 5.146", "Set MJ to NOMJremains or MJ3."),
    ("Profile manual override", "5.148 / 5.150", "Select the configured Standard or Low profile."),
    ("Toggle UKFset1 live comparison", "5.152", "Toggle the Virtual-Pump, non-Client live UKFset1 comparison; hidden elsewhere."),
    ("Run SensorAge code", "5.156", "Toggle the sensor-age code on or off."),
]

list2_rows = [
    ("MJ start", "5.158", "Start the confirmed Mounjaro cycle."),
    ("MJ restore", "5.160", "Run the confirmed Mounjaro restoration action."),
    ("Steroid start", "5.162", "Start the steroid-management sequence."),
    ("Toggle MJ Kotlin buttons", "5.164", "Show/hide the MJ buttons."),
    ("Toggle Steroid Kotlin button", "5.166", "Show/hide the steroid button."),
    ("Steroid 110 to 130", "5.168", "Run the confirmed steroid escalation to 130."),
    ("Steroid 130 to 150", "5.170", "Run the confirmed steroid escalation to 150."),
    ("Steroid 150 to 190", "5.172", "Run the confirmed steroid escalation to 190."),
    ("Steroid 190 to 250", "5.174", "Run the confirmed steroid escalation to 250."),
    ("Steroids OFF", "5.176", "Run the confirmed steroid turn-off action."),
    ("AnyDesk restart", "5.178", "Client writes an ADesk Note and uses a five-minute relay TT only if no real TT is active; otherwise Note-only transport preserves the real TT."),
    ("AnyDesk restart - local test", "5.180", "Local command test through the real receiver; no NS round-trip or relay TT."),
    ("Graph: UKF1", "No relay TT", "Local display toggle plus optional Libre slope/offset calibration."),
    ("Graph: UKF2", "No relay TT", "Local display toggle; calibration is already upstream and remains shown as on."),
    ("Graph: UKF3", "No relay TT", "Local display toggle plus optional Libre slope/offset calibration."),
    ("Graph: Graph5 panel", "No relay TT", "Local panel toggle; optional BGL-only mode hides insulin activity and the three carb series."),
]

section9 = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("9. GUI display changes"))
heading = insert_paragraph_before(section9, "d. List 1 and List 2 temporary-target command catalogue", "Heading 2")
add_bookmark(heading, "sec8_list_tt_catalogue")
insert_paragraph_before(
    section9,
    "The mmol/L numbers below are command signals, not clinical glucose targets. On a pump or Virtual build, List 1 and the clinical List 2 actions invoke the matching local handler directly and create no temporary target. On AAPSClient, the same selection creates a five-minute relay TT because preferences/actions do not sync directly; the pump receives the code, performs the action, cancels the relay target and applies its relay repeat guard. Every command still passes through its confirmation popup. List 2 graph controls are always local display settings and never use relay TTs.",
    "Normal",
)
list1_title = insert_paragraph_before(section9, "List 1 - settings and command codes", "Heading 3")
add_bookmark(list1_title, "sec8_list1_tt_codes")
add_command_table(section9, list1_rows)
insert_paragraph_before(
    section9,
    "List 1 availability note. Code 5.152 is offered only on a non-Client Virtual-Pump comparison device. Code 5.154 is intentionally unused: UKF2 graph history is maintained without switching the live dosing engine, so checking a graph option no longer changes the glucose source.",
    "Normal",
)
list2_title = insert_paragraph_before(section9, "List 2 - actions, relay codes and local graph controls", "Heading 3")
add_bookmark(list2_title, "sec8_list2_tt_codes")
add_command_table(section9, list2_rows)

# Add the new linked subsection to Appendix B Contents immediately before section 9.
contents_9 = next(
    p for p in doc.paragraphs
    if p.style.name == "Normal" and p.text.startswith("9. GUI display changes")
)
contents = insert_paragraph_before(contents_9, style="Normal")
contents.paragraph_format.left_indent = Pt(18)
contents.paragraph_format.space_after = Pt(2)
add_internal_link(contents, "d. List 1 and List 2 temporary-target command catalogue", "sec8_list_tt_catalogue")

doc.save(OUTPUT)
print(OUTPUT)
