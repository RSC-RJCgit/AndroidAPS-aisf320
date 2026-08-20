from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 subheading line spacing standardized.docx")
OUTPUT = Path("AutoISF Operations Manual mydoc Aug 21 26 carbohydrate graphs explained.docx")
ANCHOR_TEXT = "10. Factors affecting progress and hypoglycaemia outcomes"


document = Document(SOURCE)
anchor = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == ANCHOR_TEXT)


def insert_before(text: str, style: str = "Normal", bold_lead: str | None = None):
    paragraph = document.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    if style == "Normal":
        paragraph.paragraph_format.left_indent = Inches(0.5)
    else:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.keep_with_next = True
    anchor._p.addprevious(paragraph._p)
    return paragraph


insert_before("d. What the different carbohydrate graphs mean", "Heading 2")
insert_before(
    "Carbohydrate entries and COB are not absorption-rate curves. A carbohydrate treatment marker records the grams entered at a particular time. COB is the estimated grams from entered carbohydrate that remain unabsorbed and is normally selected on a secondary graph. COB therefore describes a remaining quantity, whereas the curves below describe a rate or inferred effect in grams per five minutes.",
    bold_lead="Carbohydrate entries and COB are not absorption-rate curves."
)
insert_before(
    "Carbs Absorption (C-ABS) is the empirical known-carbohydrate absorption rate. It uses autosens this5MinAbsorption, is exponentially smoothed to reduce five-minute bucket noise, and contains historical points only because future absorption has not yet been observed. It is drawn as a solid measured-data line. On graph 0 its chart-menu checkbox is the master gate, while the basal/IOB long-press display preset can temporarily show or hide it.",
    bold_lead="Carbs Absorption (C-ABS)"
)
insert_before(
    "Carb model curve is a theoretical forward curve calculated from entered carbohydrate, not a glucose-derived measurement. Each entry uses the two-compartment Gamma(2) form Ra(t) = 5 x carbs x 0.9 x k^2 x t x exp(-k x t), with k=1/90 per minute, a 90-minute peak and a six-hour cutoff. Contributions from overlapping entries are added. Because it can be calculated from known entries, it extends into the future and is drawn dashed. Its List 1 Graph2 toggle (relay code 5.138) is independent of the C-ABS checkbox.",
    bold_lead="Carb model curve"
)
insert_before(
    "UAM Carb Impact (UAMci) is the deviation-inferred carbohydrate-equivalent rate from the nearest AIV record, accepted when it lies within four minutes of the graph bucket and then exponentially smoothed. It represents glucose behaviour that looks like carbohydrate impact, including possible unannounced carbohydrate. It remains uncapped so an unusually large spike is visible diagnostically. It is historical only and is drawn as a dotted inferred-data line, distinct from the solid C-ABS line and dashed model curve.",
    bold_lead="UAM Carb Impact (UAMci)"
)
insert_before(
    "Combined Carbs (CmbC) combines known empirical absorption with only the UAM amount above that already-explained absorption. Its implemented calculation is C-ABS + max(UAMci - C-ABS, 0), which is equivalent to max(C-ABS, UAMci). This avoids counting the same known meal response twice while retaining excess impact that may represent unannounced carbohydrate. It does not include the theoretical carb-model curve. It is historical, solid and shown in its own colour.",
    bold_lead="Combined Carbs (CmbC)"
)
insert_before(
    "Scaling and interpretation. Each rate line is normalized against its own peak when placed with differently dimensioned series such as insulin activity. Equal apparent peak heights therefore support comparison of shape and timing—onset, peak and decay—but do not mean the lines have equal gram values. Read the underlying source and line style before interpreting vertical separation as a quantitative difference.",
    bold_lead="Scaling and interpretation."
)
insert_before(
    "Where the curves appear. Graph 0 can independently show C-ABS, the carb model, UAMci and CmbC. UAMci and CmbC can also be assigned to secondary graphs. Graph 5 full mode shows the carb model, UAMci and CmbC, but not the empirical C-ABS line; its BGL only checkbox hides all three of those Graph 5 carbohydrate curves. Graph 5 ignores graph 0's individual visibility selections, although the global carb-model setting must be on for model data to exist.",
    bold_lead="Where the curves appear."
)

document.save(OUTPUT)
print(OUTPUT.resolve())
