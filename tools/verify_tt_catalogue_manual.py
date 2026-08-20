from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 UKF battery and TT commands.docx")
doc = Document(path)
bookmarks = {n.get(qn("w:name")) for n in doc.element.body.xpath(".//w:bookmarkStart")}
anchors = [n.get(qn("w:anchor")) for n in doc.element.body.xpath(".//w:hyperlink[@w:anchor]")]
headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
    " | ".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows
)
required = [
    "d. List 1 and List 2 temporary-target command catalogue",
    "5.002 / 5.004", "5.152", "5.158", "5.178", "No relay TT",
    "The mmol/L numbers below are command signals, not clinical glucose targets.",
]
print({
    "tables": len(doc.tables),
    "last_table_rows": [len(t.rows) for t in doc.tables[-2:]],
    "headings": len(headings),
    "bookmarks": len(bookmarks),
    "hyperlinks": len(anchors),
    "broken_anchors": sorted(set(anchors) - bookmarks),
    "unbookmarked_headings": [p.text for p in headings if not p._p.xpath(".//w:bookmarkStart")],
    "required_missing": [s for s in required if s not in all_text],
    "back_to_toc_count": all_text.count("Back to TOC"),
})
