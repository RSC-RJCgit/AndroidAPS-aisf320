from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 progress and hypo outcomes.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 UKF and battery added.docx")
doc = Document(SOURCE)


def next_bookmark_id():
    ids = []
    for node in doc.element.body.xpath(".//w:bookmarkStart"):
        try:
            ids.append(int(node.get(qn("w:id"))))
        except (TypeError, ValueError):
            pass
    return max(ids, default=0) + 1


bookmark_id = next_bookmark_id()


def add_bookmark(paragraph, name):
    global bookmark_id
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(1, start)
    paragraph._p.append(end)
    bookmark_id += 1


def add_internal_link(paragraph, text, anchor, size_half_points=None):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    if size_half_points:
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(size_half_points))
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), str(size_half_points))
        props.extend([size, size_cs])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([props, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_paragraph_before(reference, text="", style=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    reference._element.addprevious(paragraph._element)
    return paragraph


appendix_a = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Appendix A"))

ukf = insert_paragraph_before(appendix_a, "11. UKF: formulas, roles and initial comparison assessment", "Heading 1")
add_bookmark(ukf, "sec11_ukf")

ukf_roles = insert_paragraph_before(appendix_a, "a. The UKF variants and where each is used", "Heading 2")
add_bookmark(ukf_roles, "sec11_ukf_roles")
insert_paragraph_before(
    appendix_a,
    "The labels describe different compositions of two operations - Unscented Kalman filtering and LibreSpecial exponential smoothing - rather than three interchangeable copies of one value. UKF1 is a batch, display-oriented UKF/RTS pass over the raw/noise series. Its graph can optionally apply the Libre slope and offset after filtering. UKFset1 is the live incremental counterpart that can replace the ordinary LibreSpecial path only on the non-Client Virtual-Pump comparison device; it maintains its own current state and short retrospective history. UKF2 applies the normal calibrated LibreSpecial EMA first and then refines that EMA result with UKF; its graph reads the stored post-UKF value actually used when UKF2 was live, not the earlier pre-refinement EMA. UKF3 reverses the order: it starts with UKF1, optionally applies Libre calibration, and then runs the LibreSpecial EMA. UKF3 is a display/diagnostic comparison and is not presently an independently selectable dosing source.",
    "Normal",
)

ukf_formula = insert_paragraph_before(appendix_a, "b. Formula and full explanation of each", "Heading 2")
add_bookmark(ukf_formula, "sec11_ukf_formula")
insert_paragraph_before(
    appendix_a,
    "Common UKF model. The two-element state is x = [G, r], where G is glucose in mg/dL and r is its rate in mg/dL per minute. For elapsed time dt, prediction uses G(next) = G + r x dt and r(next) = 0.98 x r. Merwe-scaled sigma points (alpha 1.0, beta 0.0, kappa 3.0) propagate the state and covariance through that model. The observation is the glucose reading itself, z = G + measurement noise. The innovation is z - predicted G; the Kalman gain determines how much of that difference corrects both glucose and rate. The estimated rate is constrained to -4 to +4 mg/dL/min. The full adaptive filter uses process covariance Q = diag(1.0, 0.40), scaled by dt/5, and adapts measurement variance R to sensor quality. The raw comparison paths deliberately use heavier fixed measurement variance R = 225 and Q = diag(1.0, 0.15), producing stronger smoothing without altering the live adaptive filter's learned state.",
    "Normal",
)
insert_paragraph_before(
    appendix_a,
    "UKF1 formula and meaning. UKF1 applies the common UKF prediction/update pass to raw readings and then, for the retrospective graph, applies Rauch-Tung-Striebel backward smoothing: x_s(k) = x_f(k) + C(k)[x_s(k+1) - x_pred(k+1)], where C is formed from the filtered covariance, transition model and next predicted covariance. This uses later points to improve earlier displayed estimates, so the retrospective UKF1 graph must not be mistaken for a value that was fully available in real time. The live UKFset1 form performs only the causal incremental prediction/update for the newest point and persists that current state. A gap over seven minutes inflates uncertainty and decays the rate; a non-positive gap or one over 60 minutes starts a fresh state. Optional graph calibration is G_cal = max(40, slope x G_UKF1 + offset), with the offset converted to mg/dL when the displayed preference is in mmol/L.",
    "Normal",
)
insert_paragraph_before(
    appendix_a,
    "LibreSpecial EMA formula. Let a be the configured smoothing factor, e the elapsed minutes, c the expected CGM interval (normally one minute; five for the relevant G7 path), and M the maximum smoothing gap. The effective factor is a_eff = min[1, a + (1-a) x {max(0,e-c)/(M-c)}^2]. The new value is S = S_previous + a_eff x (G_calibrated - S_previous); the first point uses G_calibrated directly. As the gap approaches M, a_eff approaches 1, preventing an old smoothed value from dominating after missing data.",
    "Normal",
)
insert_paragraph_before(
    appendix_a,
    "UKF2 formula and meaning. UKF2 is UKF(EMA(calibrated Libre)). Each new reading first passes through the Libre calibration and the EMA formula above. That EMA value is appended to a rolling history; a UKF/RTS pass is recomputed over up to 90 minutes/120 points, and the refined value at the current timestamp becomes the UKF2 result. Its separate post-refinement history retains the values actually returned for dosing and supplies the UKF2 graph and delta calculations. Because calibration occurs before the EMA and UKF, UKF2 has no separate graph-calibration checkbox.",
    "Normal",
)
insert_paragraph_before(
    appendix_a,
    "UKF3 formula and meaning. UKF3 is EMA(calibrate(UKF1)). Starting with UKF1's oldest-to-newest series, each point is optionally calibrated using max(40, slope x G_UKF1 + offset), then passed through the same gap-adjusted EMA. This is the opposite order from UKF2. Filtering and EMA do not commute: UKF2 lets the EMA shape what the UKF receives, whereas UKF3 lets the UKF shape what the EMA receives. Their different timing and curvature are therefore expected and are precisely what the comparison graph is intended to reveal.",
    "Normal",
)

ukf_progress = insert_paragraph_before(appendix_a, "c. Progress of the initial UKF-check assessment and earliest-change question", "Heading 2")
add_bookmark(ukf_progress, "sec11_ukf_progress")
insert_paragraph_before(
    appendix_a,
    "The initial observation was that the UKF1 graph appeared to react to a fall before LibreSpecial. The current code treats that as a hypothesis requiring repeated measurement, not as proof that UKF1 should replace the live source. On the non-Client Virtual-Pump comparison device only, it keeps the alternative histories warm and compares the loop value, UKF2, live-style UKFset1, a shadow LibreSpecial value, retrospective UKF1 and UKF3. Every five minutes it records each type's 5-, 15- and 30-minute deltas plus fitted acceleration, positive/negative parabola deltas, fit window and correlation quality. It separately warns when a type's meaningful acceleration sign disagrees with the live loop.",
    "Normal",
)
insert_paragraph_before(
    appendix_a,
    "Earliest-change tracking defines a genuine falling crossing as delta5 at or below -0.2 mmol/L per five minutes. The loop crossing is checked every cycle; the five alternative-type comparisons are normally sampled by their five-minute diagnostic blocks. For each falling episode, the first crossing time for every type is retained in memory and later crossings are reported as seconds behind the first. Recovery above the threshold re-arms that type for the next episode. A dedicated LibreSpecial-versus-UKFset1 race removes interference from the other variants: the leader writes LibreLd or Set1Ld to CarePortal, while both lead/lag times remain in the log. The UKFcheck export gathers these metrics and race records during manual AIV export and the normal six-hour export cycle.",
    "Normal",
)
insert_paragraph_before(
    appendix_a,
    "Use beyond the assessment. At present the earliest signal is diagnostic only: it does not automatically change the live glucose source, dose, temporary target or automation. A single early crossing may be genuine lead time, retrospective RTS advantage, or a false lead caused by raw noise. Permanent or external use should therefore wait for enough matched episodes to compare median lead time, missed declines, false alarms, low-point error and recovery lag using only values that were causally available at that moment. In particular, retrospective UKF1/UKF3 graph points must not be credited as real-time warning unless the same advantage is reproduced by the causal UKFset1 history. If those findings remain favorable, the live UKFset1 toggle is the appropriate candidate for a controlled trial; the diagnostic should not itself select whichever line happened to cross first in each episode.",
    "Normal",
)

battery = insert_paragraph_before(appendix_a, "12. Battery 1% automation and profile recovery", "Heading 1")
add_bookmark(battery, "sec12_battery")

battery_low = insert_paragraph_before(appendix_a, "a. Battery1pc: action at 1% or below", "Heading 2")
add_bookmark(battery_low, "sec12_battery_low")
insert_paragraph_before(
    appendix_a,
    "Battery1pc is checked with a 20-minute floor. It can fire only when the active profile percentage is 100%, Automation State Profile is PP130, C100 or AllOK, the receiving phone reports battery at or below 1%, and the active pump is not Virtual Pump. It deliberately has no pod/cannula-age condition: critical phone battery applies equally to patch and tubed live-pump systems. When it fires, it switches immediately to the named Current Profile50, raises the urgent Batt1% notification, adds the main-graph announcement Batt1% and CarePortal note Bt<1%, sends LowBattery by the ordinary SMS route and to the configured battery-alert numbers, and records the run. The profile switch makes the 100% precondition false on the next loop, providing an additional self-guard against repeated action.",
    "Normal",
)

battery_recovery = insert_paragraph_before(appendix_a, "b. BatteryOver1pc: restoration after recovery", "Heading 2")
add_bookmark(battery_recovery, "sec12_battery_recovery")
insert_paragraph_before(
    appendix_a,
    "BatteryOver1pc is checked with a five-minute floor. It requires battery above 1% and the exact current profile name Current Profile50. The profile-name test is intentional: a hypo-created 50% percentage on ProfileReal is not the battery profile and must not be cancelled merely because the battery recovered. The older Profile=Batt1% state test could never succeed because Battery1pc did not set that state. On recovery, the automation switches to the configured Standard Profile, sends AllOK Batt, sets Automation State Profile to AllOK, writes CarePortal note bat>1 and records the recovery run. There is no separate user confirmation; progression is Battery1pc reduction followed by BatteryOver1pc restoration once these exact conditions are observed.",
    "Normal",
)

# Add linked items to the existing numbered Introduction.
first_section = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("1. Requirements"))
first_index = next(i for i, p in enumerate(doc.paragraphs) if p._element is first_section._element)
intro_numbered = [p for p in doc.paragraphs[:first_index] if p._p.xpath("./w:pPr/w:numPr")]
if not intro_numbered:
    raise RuntimeError("Introduction numbered list not found")
num_pr = deepcopy(intro_numbered[0]._p.xpath("./w:pPr/w:numPr")[0])
for text, anchor in [
    ("UKF formulas, roles and initial comparison assessment.", "sec11_ukf"),
    ("Battery 1% automation and profile recovery.", "sec12_battery"),
]:
    p = doc.add_paragraph(style="Normal")
    p._p.get_or_add_pPr().append(deepcopy(num_pr))
    p.paragraph_format.space_after = Pt(3)
    add_internal_link(p, text, anchor, 24)
    first_section._element.addprevious(p._element)

# Extend Appendix B's linked contents before its Appendix A entry.
contents_appendix_a = next(
    p for p in doc.paragraphs
    if p.text.startswith("Appendix A") and p._element is not appendix_a._element and p.style.name == "Normal"
)
contents_entries = [
    ("11. UKF: formulas, roles and initial comparison assessment", "sec11_ukf", False),
    ("a. The UKF variants and where each is used", "sec11_ukf_roles", True),
    ("b. Formula and full explanation of each", "sec11_ukf_formula", True),
    ("c. Progress of the initial UKF-check assessment and earliest-change question", "sec11_ukf_progress", True),
    ("12. Battery 1% automation and profile recovery", "sec12_battery", False),
    ("a. Battery1pc: action at 1% or below", "sec12_battery_low", True),
    ("b. BatteryOver1pc: restoration after recovery", "sec12_battery_recovery", True),
]
for text, anchor, indented in contents_entries:
    p = insert_paragraph_before(contents_appendix_a, style="Normal")
    if indented:
        p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    add_internal_link(p, text, anchor)

doc.save(OUTPUT)
print(OUTPUT)
