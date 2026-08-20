from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 underlined indented headings.docx")
doc = Document(path)
bookmarks = {n.get(qn("w:name")) for n in doc.element.body.xpath(".//w:bookmarkStart")}
anchors = [n.get(qn("w:anchor")) for n in doc.element.body.xpath(".//w:hyperlink[@w:anchor]")]
h1 = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
h2 = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
h3 = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
print({
    "tables": len(doc.tables),
    "hyperlinks": len(anchors),
    "broken_anchors": sorted(set(anchors) - bookmarks),
    "heading1_not_underlined": [p.text for p in h1 if any(r.text and r.font.underline is not True for r in p.runs)],
    "heading2_wrong_indent_pt": [(p.text, p.paragraph_format.left_indent.pt if p.paragraph_format.left_indent else None) for p in h2 if not p.paragraph_format.left_indent or abs(p.paragraph_format.left_indent.pt - 18) > 0.1],
    "heading3_wrong_indent_pt": [(p.text, p.paragraph_format.left_indent.pt if p.paragraph_format.left_indent else None) for p in h3 if not p.paragraph_format.left_indent or abs(p.paragraph_format.left_indent.pt - 36) > 0.1],
    "unbookmarked_headings": [p.text for p in h1 + h2 + h3 if not p._p.xpath(".//w:bookmarkStart")],
})
