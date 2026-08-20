from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn


SOURCE = Path(r"C:\winword\aaa\AutoISF Idiosyncrasies and Operations Manual mydoc Aug 21 26 all boosts.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_Operations_Manual_mydoc_Aug21_2026_section_C.docx")


def remove(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def strip_navigation(doc):
    body = doc._element.body
    title = next(p for p in doc.paragraphs if p.style.name == "Title")
    for child in list(body):
        if child is title._element:
            break
        remove(child)
    for hyperlink in list(body.iter(qn("w:hyperlink"))):
        visible = "".join(t.text or "" for t in hyperlink.iter(qn("w:t")))
        if visible.strip() == "Back to TOC":
            remove(hyperlink)
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for bookmark in list(body.iter(qn(tag))):
            remove(bookmark)
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            style = p.style.name
            p.text = p.text.strip()
            p.style = doc.styles[style]


def body(doc, label, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(label).bold = True
    p.add_run(text)
    return p


doc = Document(SOURCE)
strip_navigation(doc)
appendix_anchor = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Appendix A"))

section_start = doc.add_heading("8. Manual controls and AIV history", level=1)

doc.add_heading("a. Bolus calculator: screen, calculation and follow-up dosing", level=2)
body(doc, "Initial screen and live calculation. ",
     "The Wizard recalculates after each entry or checkbox change. The screen can include BG and target, carbs, protein, fat, correction (units or percentage), carb time, profile, BG trend, COB, bolus and basal IOB, superbolus, the session-only maximum-bolus override, and the calculation breakdown. Selecting COB also selects IOB; clearing IOB clears COB. The displayed total is the constrained, pump-step-rounded proposal. A calculated 0 U remains confirmable so that the decision can be recorded. The maximum-bolus override affects this Wizard session only and is restored on every exit path.")
body(doc, "Profile-percentage behavior. ",
     "At an active 50% profile, the normal immediate calculation is reduced through the profile's effective IC and the automatic delayed-bolus mechanism may later check whether the remaining gap is justified. A separate LowBG=50recent state, or BG below 5.0 mmol/L and not rising, halves only the carbs-driven component when the active profile is not already 50%; the two reductions cannot stack. The manual carb-split row is displayed only when the feature is enabled, the active profile is exactly 100%, the proposed dose exceeds the maximum allowed bolus, and that maximum is positive. The delivery engine accepts an already-selected split at 100% or higher, but the present screen does not offer the row at a profile above 100%. Protein/fat projections can still be shown independently; their delayed delivery requires a profile of at least 100%.")
body(doc, "Component calculation. ",
     "The immediate proposal combines BG correction against the selected profile or temporary target, a 15-minute trend contribution, carbs divided by IC, included COB divided by IC, the negative of included bolus/basal IOB, a manual correction, and any selected superbolus. Protein and fat are not included in the immediate total: protein is converted as protein/25 divided by IC and fat as fat/11 divided by IC, then held as separate delayed-dose projections. The Wizard percentage is applied to the non-protein/fat immediate total before pump-step rounding and the normal bolus-constraint pipeline.")
body(doc, "Fast-rise projected-HP protection. ",
     "The current BolusWizard path, including Quick Wizard calculations using that engine, applies the retrospective guard only when projected HP is at most 6.5 mmol/L, current BG is below 10.0 mmol/L, both current and short-average rise are at least 0.6 mmol/L per five minutes, positive IOB is at least 2.0 U, and the proposed dose is at least 0.75 U. It first removes the calculated COB-insulin contribution and then reduces the remaining already percentage-adjusted proposal to 75%. The changed dose is recalculated and shown as the proposal; the saved calculator result records original dose, adjusted dose, projected HP and COB insulin removed. This is automatic, not a second optional action in the confirmation dialog.")
body(doc, "Manual split-bolus projection and first delivery. ",
     "When the split row is available it starts checked and uses the configured interval (editable from 1 to 60 minutes). The screen shows the planned number and size of parts: the maximum allowed bolus is the first part, with the residual projected at the selected interval. After confirmation, the first constrained part is delivered immediately. The residual is scheduled only once, and a CarePortal scheduling marker records the projected future total. SMBs remain allowed during this manual split sequence.")
body(doc, "Manual split-bolus progress and recalculation. ",
     "Before each later part the code verifies that the user has not stopped follow-up dosing, no newer bolus/carbs entry has superseded the schedule, profile remains at least 100%, the pump is available, and superbolus is not active. It obtains fresh glucose. Missing data is retried after two minutes. BG below 6.0 mmol/L is unsafe; from 6.0 to below 8.0 mmol/L, a delta or short-average delta at or below -0.05 mmol/L is also unsafe. The first two consecutive unsafe checks defer by a full split interval; the third cancels the residual. The complete soft-retry window is limited to 60 minutes. Each part is reduced by any positive rise in live IOB since the preceding post-dose baseline and rounded to the pump step. A result at or below zero is skipped and retried at the next interval rather than increasing the dose. Success schedules the next residual; failure raises the normal delivery alarm. Cancellation and calculated-zero outcomes are written to treatment/history markers.")
body(doc, "Automatic delayed bolus at 50% or recent-50. ",
     "This is distinct from the manual split. After the initial Wizard/Quick Wizard bolus, the worker checks every 10 minutes from +10 through +80 minutes while temporarily blocking SMBs. It requires glucose data no more than five minutes old, BG above 4.5 mmol/L, current delta above +0.10 mmol/L, short-average delta above +0.20 mmol/L, and long-average delta above +0.05 mmol/L. When the criteria first pass, it calculates the frozen original full-required amount minus the original delivered dose, then delivers 90% of that gap if confirmation occurs by 20 minutes or 50% if it occurs later. It does not recompute the original IOB snapshot. Db10, Db20 and subsequent check markers show dose, wait or end; SMBs are unblocked on delivery, cancellation or final expiry.")
body(doc, "Protein and fat delayed doses. ",
     "Protein is scheduled once at +120 minutes and fat once at +180 minutes, independently of the manual carb split. At delivery each requires the schedule still to be current, profile at least 100%, pump available, no superbolus, fresh BG at least 7.0 mmol/L, and both delta measures above -0.05 mmol/L. A positive rise in IOB since the immediate post-bolus baseline is deducted. If the remainder is zero or negative, a 0 U history entry plus cancellation note is recorded; otherwise the rounded remainder is delivered. These fixed-time doses do not retry after a failed BG gate.")
body(doc, "Confirmation and visible history. ",
     "Pressing OK performs one final recalculation, copies the split checkbox/interval into the BolusWizard, opens the standard confirmation/execution flow, and prevents a second OK tap from duplicating it. The BolusCalculatorResult keeps the calculation inputs and notes, including any carb-split projection, protein/fat timing and HP-safety adjustment. CarePortal notes and treatment entries then show scheduling, checks, delivered parts, zero-dose outcomes and cancellation reasons; these entries are the durable status trail after the original dialog closes.")

doc.add_heading("b. Lists 1 and 2: mechanism, current status and confirmation", level=2)
body(doc, "List 1 entry and purpose. ",
     "Double-tap the IOB area to open List 1. On a pump/Virtual build its title is Direct AutoISF settings and a confirmed selection invokes the matching local AutoISF handler without creating a temporary target. On AAPSClient its title is Settings - relay to pump; preferences are not changed locally, so confirmation creates the matching five-minute custom temporary-target code for the pump to receive, apply and cancel.")
body(doc, "List 1 rows and current status. ",
     "The outer list shows setting names rather than raw relay numbers. It includes paired adjustments for SMB-delivery values, normal/high ISF weights, Libre slope/offset, SMB offset, Wizard percentage, mild-boost ratio, insulin peak, AutoISF limits and eight time-of-day offsets; single actions/toggles include sensor sensitivity, all boost automations, clean graph, Graph2/Graph5, cloud logs, UKFset1 where permitted and SensorAge; named paired controls cover MJ state and Standard/Low profile. Opening a row shows Current on pump/Virtual. A client shows Pump current from the latest mirrored device-status settings snapshot, with its age, or unavailable if the key has not arrived. Dual-key rows show both values. Stateless commands show an Effect description instead of pretending to have a current state.")
body(doc, "List 1 confirmation. ",
     "A single/toggle row opens its description/current-value popup and requires OK. A stepped row opens mutually exclusive decrease/increase (or named-state) checkboxes labelled with the real change; the selection is not applied until OK. OK with neither direction selected is a no-op. Cancel, back, or outside-dismiss from an inner popup returns to List 1 so another row can be selected; Cancel on the outer list closes to Overview. Client relay commands retain the pump-side repeat guard; direct pump/Virtual actions do not use that relay guard.")
body(doc, "List 2 entry and actions. ",
     "Double-tap the basal-rate area to open List 2. Its direct-action rows cover MJ start/restore, steroid start and escalation/turn-off, MJ and steroid button toggles, and the appropriate AnyDesk restart choices. On pump/Virtual, confirmed actions are dispatched locally. On AAPSClient, most are carried by their five-minute relay temporary-target code. AnyDesk is different: it always writes an ADesk Note; it adds a relay TT only when no real temporary target is active, preserving an existing target. The local-test AnyDesk row queues the real receiving-side path and records its own diagnostic note.")
body(doc, "List 2 graph controls and current status. ",
     "The final rows control UKF1, UKF2, UKF3 and the Graph5 panel. Selecting one opens checkboxes showing the current local display state: Graph on plus its calibration option; UKF2 calibration is shown checked but disabled because calibration is already applied upstream. Graph5's second checkbox means BGL only. These settings are display-local and are written only after OK. Direct-action rows provide a confirmation question but do not all expose a universal live state; their resulting automation state, profile, target or note is the authoritative status after execution.")
body(doc, "List 2 confirmation and return behavior. ",
     "Every direct action requires the confirmation popup. Cancel returns to List 2. For a graph row, OK saves the selected checkboxes and refreshes Overview; Cancel, back or outside-dismiss returns to List 2. The outer Cancel closes the list. Client List 2 is therefore a command/relay interface, while its graph rows remain immediate local display settings.")

doc.add_heading("c. AIV table: popup population and backup/export behavior", level=2)
body(doc, "Opening and background population. ",
     "Long-pressing the ISF/history control opens the AIV history popup. On first creation it reads the preceding six hours of AIV records in descending time order, APS results, step counts, SMB boluses and CarePortal notes. It also reads raw glucose with a 20-minute lead-in for delta calculations and computes the UKF3 comparison series. The work is assembled from the local database; it is not populated by parsing the displayed graph. A frozen Time column remains aligned with the horizontally scrolling header/body. SMB only rebuilds the visible table from records whose delivered-SMB value is positive, and All restores every loaded record; filtering does not change the full-record lookback used for derived values. Rotation preserves the filter and does not start another export.")
body(doc, "What the popup represents. ",
     "Each row joins the nearest available context to an AIV timestamp and presents grouped glucose/raw-delta, ISF-factor, insulin/carb/HP, request/delivery and activity/status fields. A missing six-hour history produces an explicit no-data row. Values such as COBt, IOB change, HP/HP2/HP3 and UKF-derived columns are calculated by the exporter/table helpers from the loaded windows; they should therefore be interpreted with their source timing and availability, not as independent sensor samples.")
body(doc, "Manual/on-open export. ",
     "The first opening of the popup starts an off-UI-thread export immediately. It writes three current AIV files (CSV, text and settings), refreshes the combined export, records success or failure in logs and an AVLs/AVLf CarePortal marker, then attempts cloud upload when cloud storage is configured. After the AIV upload callback it can invoke the same log-send route without exporting AIV a second time. Opening the popup is therefore the manual refresh mechanism, although the file work itself continues in the background and the table remains usable.")
body(doc, "Automatic backup/export. ",
     "KeepAlive checks the AIV export timestamp and normally performs the same last-six-hours three-file export every six hours, also rebuilding the combined file and its dated per-device copy. The files are always written locally first. When cloud storage is active they are additionally uploaded to the configured AIV cloud path, scoped by General patient name when supplied. If automatic cloud log export is enabled and due, KeepAlive reserves that six-hour slot before launching the asynchronous AIV-plus-log sequence so overlapping KeepAlive runs do not duplicate it. Automatic operation is silent: success/failure is logged rather than shown as a popup.")
body(doc, "Backup scope and recovery meaning. ",
     "These AIV files are history/diagnostic exports, not a restorable database backup. Manual popup-open and automatic six-hour runs create equivalent current exports and may upload copies; neither reimports AIV rows into the application. The combined/dated files preserve longer troubleshooting continuity, while each popup view itself remains a six-hour database window. A failed cloud upload does not remove the local files, and absence of cloud configuration simply leaves the local export as the backup copy.")

# Keep numbered operating sections before the installation appendix.
moving = False
for element in list(doc._element.body):
    if element is section_start._element:
        moving = True
    if moving and element.tag != qn("w:sectPr"):
        appendix_anchor._element.addprevious(element)

doc.save(OUTPUT)
print(OUTPUT)
