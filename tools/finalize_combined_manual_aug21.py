from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path("AutoISF Operations Manual mydoc Aug 21 26 complete with Settings and Automations.docx")
OUTPUT = Path("AutoISF Operations Manual mydoc Aug 21 26 complete Settings Automations final.docx")


def clean_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        if "Back to TOC" in run.text:
            run.text = run.text.replace("Back to TOC", "").rstrip()


document = Document(SOURCE)
for paragraph in document.paragraphs:
    clean_paragraph(paragraph)
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                clean_paragraph(paragraph)

# Hyperlink-contained runs are not exposed by older python-docx paragraph.runs.
# Clean every Word text node as a final deterministic pass.
for text_node in document.element.body.iter(qn("w:t")):
    if text_node.text and "Back to TOC" in text_node.text:
        text_node.text = text_node.text.replace("Back to TOC", "").rstrip()

document.save(OUTPUT)
print(OUTPUT.resolve())
