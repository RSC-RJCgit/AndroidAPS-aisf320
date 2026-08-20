from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 UKF and battery added.docx")
doc = Document(path)
bookmarks = {n.get(qn("w:name")) for n in doc.element.body.xpath(".//w:bookmarkStart")}
anchors = [n.get(qn("w:anchor")) for n in doc.element.body.xpath(".//w:hyperlink[@w:anchor]")]
headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
text = "\n".join(p.text for p in doc.paragraphs)
required = [
    "11. UKF: formulas, roles and initial comparison assessment",
    "b. Formula and full explanation of each",
    "c. Progress of the initial UKF-check assessment and earliest-change question",
    "12. Battery 1% automation and profile recovery",
    "Battery1pc is checked with a 20-minute floor",
    "BatteryOver1pc is checked with a five-minute floor",
    "UKF formulas, roles and initial comparison assessment.",
    "Battery 1% automation and profile recovery.",
]
print({
    "tables": len(doc.tables),
    "headings": len(headings),
    "bookmarks": len(bookmarks),
    "hyperlinks": len(anchors),
    "broken_anchors": sorted(set(anchors) - bookmarks),
    "unbookmarked_headings": [p.text for p in headings if not p._p.xpath(".//w:bookmarkStart")],
    "required_missing": [s for s in required if s not in text],
    "back_to_toc_count": text.count("Back to TOC"),
})
