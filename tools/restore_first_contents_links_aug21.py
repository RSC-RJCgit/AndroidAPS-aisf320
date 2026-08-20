from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 additions moved from Contents to main.docx")
OUTPUT = Path("AutoISF Operations Manual mydoc Aug 21 26 first Contents links restored.docx")

LINKS = (
    ("d. Duration (dura) modifications to the AutoISF algorithm", "sec_dura_mod"),
    ("e. TDD factor / sensitivity toggles", "sec_tdd_factor"),
    ("f. The overnight safety-guard architecture", "sec_overnight_guard"),
    ("d. What the different carbohydrate graphs mean", "sec_carb_graphs"),
)


document = Document(SOURCE)


def matches(text: str, style: str):
    return [p for p in document.paragraphs if p.text.strip() == text and p.style.name == style]


def next_bookmark_id() -> int:
    ids = []
    for start in document.element.body.iter(qn("w:bookmarkStart")):
        value = start.get(qn("w:id"))
        if value and value.isdigit():
            ids.append(int(value))
    return max(ids, default=0) + 1


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    children = list(paragraph._p)
    insert_at = 1 if children and children[0].tag == qn("w:pPr") else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def nearby_link_run_properties(paragraph):
    body_paragraphs = document.paragraphs
    index = next(i for i, p in enumerate(body_paragraphs) if p._p is paragraph._p)
    for prior in reversed(body_paragraphs[max(0, index - 8):index]):
        hyperlink = prior._p.find(qn("w:hyperlink"))
        if hyperlink is None:
            continue
        run = hyperlink.find(qn("w:r"))
        if run is not None:
            rpr = run.find(qn("w:rPr"))
            if rpr is not None:
                return deepcopy(rpr)
    return None


def replace_with_internal_link(paragraph, text: str, anchor: str):
    rpr = nearby_link_run_properties(paragraph)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = "  " + text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


bookmark_id = next_bookmark_id()
for text, anchor_name in LINKS:
    contents = matches(text, "Normal")
    targets = matches(text, "Heading 2")
    if not contents or len(targets) != 1:
        raise RuntimeError(f"Cannot resolve link {text!r}: contents={len(contents)}, targets={len(targets)}")
    first_contents = min(contents, key=lambda p: next(i for i, q in enumerate(document.paragraphs) if q._p is p._p))
    target = targets[0]
    add_bookmark(target, anchor_name, bookmark_id)
    replace_with_internal_link(first_contents, text, anchor_name)
    bookmark_id += 1

# Preserve the original bookmark names still used by the later duplicate
# Contents list for the two headings that were renumbered from d/e to e/f.
for text, legacy_name in (
    ("e. TDD factor / sensitivity toggles", "sec014"),
    ("f. The overnight safety-guard architecture", "sec015"),
):
    add_bookmark(matches(text, "Heading 2")[0], legacy_name, bookmark_id)
    bookmark_id += 1

document.save(OUTPUT)
print(OUTPUT.resolve())
