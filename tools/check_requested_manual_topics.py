from pathlib import Path
from docx import Document

files = [
    Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 progress and hypo outcomes.docx"),
    Path(r"C:\winword\aaa\AutoISF Settings Reference mydoc Aug 21 26 current column.docx"),
    Path(r"C:\winword\aaa\AutoISF_Automations_List5 mydoc.docx"),
]
terms = ["UKF", "Libre", "formula", "battery", "1%", "BatteryOver1pc", "temporary target", "temp target", "List 1", "List 2"]
for path in files:
    doc = Document(path)
    blocks = [(f"P{i}", p.text) for i, p in enumerate(doc.paragraphs)]
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            blocks.append((f"T{ti}R{ri}", " | ".join(c.text for c in row.cells)))
    print("\nFILE", path.name, "paragraphs", len(doc.paragraphs), "tables", len(doc.tables))
    for term in terms:
        hits = [(where, text.replace("\n", " ")[:350]) for where, text in blocks if term.lower() in text.lower()]
        print(" TERM", repr(term), "HITS", len(hits))
        for where, text in hits[:5]:
            print("  ", where, repr(text))
