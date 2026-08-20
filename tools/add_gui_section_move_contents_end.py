from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn


SOURCE = Path(r"C:\winword\aaa\AutoISF Idiosyncrasies and Operations Manual mydoc Aug 21 26 intro automation additions.docx")
INTERMEDIATE = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_mydoc_gui_intermediate.docx")
NAV_INPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_mydoc_gui_nav.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_Operations_Manual_mydoc_Aug21_2026_GUI_contents_end.docx")


def remove(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def strip_navigation(doc):
    body = doc._element.body
    contents = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text == "Contents")
    first_section = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("1. Requirements"))
    removing = False
    for child in list(body):
        if child is contents._element:
            removing = True
        if child is first_section._element:
            removing = False
        if removing:
            remove(child)
    for hyperlink in list(body.iter(qn("w:hyperlink"))):
        visible = "".join(t.text or "" for t in hyperlink.iter(qn("w:t")))
        if visible.strip() == "Back to TOC":
            remove(hyperlink)
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for bookmark in list(body.iter(qn(tag))):
            remove(bookmark)
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            style = p.style.name
            p.text = p.text.replace(" Back to TOC", "").strip()
            p.style = doc.styles[style]


def labeled(doc, label, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(label).bold = True
    p.add_run(text)
    return p


doc = Document(SOURCE)
strip_navigation(doc)
first_section = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("1. Requirements"))
last_intro = next(p for p in doc.paragraphs if p.text.startswith("Bolus-calculator HP protection"))
gui_intro = doc.add_paragraph("GUI display changes.", style="Normal")
num_pr = last_intro._p.xpath("./w:pPr/w:numPr")
if num_pr:
    gui_intro._p.get_or_add_pPr().append(deepcopy(num_pr[0]))
first_section._element.addprevious(gui_intro._element)

appendix = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Appendix A"))
new_elements = []

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    new_elements.append(p._element)

def add_body(label, text):
    p = labeled(doc, label, text)
    new_elements.append(p._element)

add_heading("9. GUI display changes", 1)
add_heading("a. Graph types and the triggers for each", 2)
add_body("Main glucose graph (graph 0). ",
         "This graph is always the primary time-aligned display. Its chart-menu switches select predictions, treatments/therapy events, activity, empirical carb absorption, UAM carb impact, combined carbs, BG parabola, raw BG, smoothed raw/UKF traces and basal data. The separate Graph2 carb-model preference controls the theoretical carb-model curve independently of the empirical absorption switch. Targets, running-mode markers, the current-time line and the principal AutoISF annotations are added automatically.")
add_body("Secondary graphs (graphs 1-4). ",
         "The Overview chart menu creates and labels a secondary panel when at least one eligible series is assigned to it. Available series include absolute IOB, IOB, COB, IOB threshold, deviation, -BGI, autosensitivity ratio, variable sensitivity, deviation slope, heart rate, steps, final/acce/BG/postprandial/duration AutoISF factors, raw BG, smoothed raw/UKF data, UAM carb impact and combined carbs. The first active type normally establishes the scale; compatible IOB/threshold, deviation/-BGI and multiple AutoISF-factor selections share aligned scales.")
add_body("Graph 5. ",
         "Graph 5 is a dedicated main-sized comparison panel, not another chart-menu secondary graph. It is enabled by the Graph5 preference, the List 1 Graph5 toggle or List 2's Graph: Graph5 panel control. It deliberately ignores graph 0's individual visibility switches and always draws its own BGL comparison set. List 2 also provides BGL only, which hides insulin activity and the three carb-related curves while retaining BGL, basal and annotation rows.")
add_body("Gesture and remote triggers. ",
         "A basal-area long-press cycles the main graph's display preset and controls the quick visibility of empirical carb line 1. An IOB-area long-press toggles SMB labels, resets the basal display preset to normal and forces carb line 1 on. The Clean graph command (List 1/relay code 5.042) applies the no-SMB-label/plain BGL preset once and then clears its request flag. Double-tapping the basal area opens List 2, where UKF1/UKF2/UKF3 and Graph5 checkboxes change local graph display settings; UKF2's calibration box is informational because calibration is already applied upstream.")

add_heading("b. Colour changes driven by AutoISF input", 2)
add_body("Dominant AutoISF colour. ",
         "For each AIV point the display compares how far the acceleration, BG, postprandial and duration ISF factors lie from neutral 1.0. The factor with the greatest deviation supplies its theme colour. That dominant colour can tint the corresponding glucose dot, SMB delivery marker and the AutoISF temp-basal overlay, so the graph shows which AutoISF input was most influential at that time rather than merely whether the final factor was above or below 1.0.")
add_body("Safety and fallback precedence. ",
         "A low glucose dot always uses the low-glucose colour and is never masked by an AutoISF tint. Without a usable dominant-factor override, high points use the normal high colour and other points use the original/default BG colour. Prediction dots retain their IOB, COB, UAM or zero-temp prediction colours. When the plain/clean preset is active, BG dots revert to the conventional low, in-range and high colours and the ISF-coloured temp-basal overlay is suppressed.")
add_body("Other fixed colour conventions. ",
         "Raw BG is split into colour-banded segments rather than shown as one undifferentiated trace. UKF1, UKF2 and UKF3 retain distinct comparison-line styling. The AIV/ISF index row uses separate colours for final, acceleration, BG, postprandial and duration factors; an SMB value of zero falls back to the ordinary insulin colour, while a delivered SMB can use the dominant AutoISF colour.")

add_heading("c. Information shown on each graph", 2)
add_body("Graph 0 information. ",
         "The main graph can show measured and bucketed glucose, prediction traces, the shaded target range, bolus/carbohydrate treatments, effective-profile switches, therapy events, basals, temporary target, running modes and the now line. Optional analytical curves include insulin activity, empirical absorption, theoretical carb model, UAM and combined carbs, parabola/prediction, raw Libre and UKF-smoothed comparisons. Fixed annotations show the AAPS one-minute delta, UKF/Libre one-minute delta, HP/hypoglycaemia projection and IOB-peak information. Tapping a BG point or treatment displays its detail label.")
add_body("Graph 1 information. ",
         "In addition to its user-selected secondary series, graph 1 carries the target-offset/last-duration-taper row, stacked steps, the extra DR/AW/LS status row, noisy/raw-delta information, and the colour-segmented index row showing final, acceleration, BG, postprandial, duration and SMB values. These rows are fixed additions and therefore may appear even when the graph's principal selectable curve is something else.")
add_body("Graph 2 information. ",
         "Graph 2 shows its selected secondary series and the per-SMB label/arrow layer. SMB label visibility follows the IOB long-press/Clean graph state, so the underlying numeric curve can remain while dose labels are hidden.")
add_body("Graph 3 information. ",
         "Graph 3 shows its selected secondary series plus the total-SMB-per-stack labels at the base of the panel. The older pp/acce/duration annotation row no longer belongs here; it has moved to Graph 5.")
add_body("Graph 4 information. ",
         "Graph 4 shows its selected secondary series. When the main Treatments switch is enabled it also shows CarePortal note text and matching arrowheads in the upper part of the panel, separating diagnostic notes from the SMB stack labels now placed on graphs 2 and 3.")
add_body("Graph 5 information. ",
         "Graph 5 provides a plain comparison picture: measured/bucketed glucose and predictions, shaded range, BG parabola, raw BG, all three UKF comparison lines, basals, target, running modes and now line. In full mode it also shows insulin activity, theoretical carb model, UAM carb impact and combined carbs. Its fixed annotation rows show target offset/last duration taper and pp/acce/duration weights. It intentionally omits treatments, therapy-event/note text, delta labels and the hypo-prediction row so marker text does not obscure the comparison traces.")

for element in new_elements:
    appendix._element.addprevious(element)

doc.save(INTERMEDIATE)
print(INTERMEDIATE)


def move_contents_to_end():
    nav = Document(NAV_INPUT)
    title = next(p for p in nav.paragraphs if p.style.name == "Title")
    toc_elements = []
    for child in list(nav._element.body):
        if child is title._element:
            break
        toc_elements.append(child)
    sect_pr = nav._element.body.find(qn("w:sectPr"))
    for child in toc_elements:
        sect_pr.addprevious(child)
    nav.save(OUTPUT)
    print(OUTPUT)
