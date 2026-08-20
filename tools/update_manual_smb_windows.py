from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 first Contents links restored.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 SMB time windows.docx")


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if paragraph._p.pPr is not None:
        inserted._p.insert(0, deepcopy(paragraph._p.pPr))
    inserted.add_run(text)
    return inserted


doc = Document(SOURCE)

smb_anchor = next(
    p for p in doc.paragraphs
    if p.text.startswith("• maxIOB threshold percent for SMB")
)
insert_after(
    smb_anchor,
    "Overnight SMB delivery reduction. From 00:00 inclusive until 08:00 exclusive "
    "(00:00–07:59), AutoISF reduces the effective SMB delivery ratio by 0.03 before "
    "calculating the SMB: effective ratio = max(0, configured smb_delivery_ratio − 0.03). "
    "For example, 0.50 becomes 0.47. This is a ratio reduction, not a fixed subtraction "
    "of 0.03 U from the resulting bolus, so its unit effect scales with the calculated "
    "insulin requirement. The configured ratio resumes at 08:00. The ordinary maximum-"
    "bolus, IOB, recent-delivery, interval, rounding and other shared SMB protections still "
    "apply, and the adjustment is recorded in the APS reason/debug output."
)

duration = next(
    p for p in doc.paragraphs
    if p.text.startswith("Duration: Tier 3 applies to one calculation cycle only.")
)
duration.text = (
    "Duration and permitted hours: Tier 3 applies to one calculation cycle only and may "
    "be initiated only from 09:00 inclusive until 21:00 exclusive (09:00–20:59), using "
    "the device's local time. Outside that window Tier 3 does not boost the candidate, "
    "although the ordinary SMB pathway may still calculate and deliver an SMB under its "
    "normal rules. Tier 3 has no persistent boost duration and no separate two-minute "
    "Tier 3 throttle. The normal SMB delivery interval still controls when another SMB may "
    "be delivered. A concurrent 30- or 60-minute low temporary basal can be produced by "
    "the ordinary SMB algorithm, but that is not a Tier 3 duration."
)

doc.save(OUTPUT)
print(OUTPUT)
