from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\winword\aaa\AutoISF Idiosyncrasies and Operations Manual mydoc Aug 21 26 complete controls and AIV.docx")
INTERMEDIATE = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_mydoc_intro_intermediate.docx")
NAV_INPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_mydoc_intro_nav.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_Operations_Manual_mydoc_Aug21_2026_introduction_then_contents.docx")


def remove(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def strip_generated_navigation(doc):
    body = doc._element.body
    title = next(p for p in doc.paragraphs if p.style.name == "Title")
    for child in list(body):
        if child is title._element:
            break
        remove(child)
    for hyperlink in list(body.iter(qn("w:hyperlink"))):
        visible = "".join(t.text or "" for t in hyperlink.iter(qn("w:t")))
        if visible.strip() == "Back to TOC":
            remove(hyperlink)
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for bookmark in list(body.iter(qn(tag))):
            remove(bookmark)
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            style = p.style.name
            p.text = p.text.replace(" Back to TOC", "").strip()
            p.style = doc.styles[style]


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720"); tabs.append(tab)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref); numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId"); num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num]); p_pr.append(num_pr)


doc = Document(SOURCE)
strip_generated_navigation(doc)
first_section = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("1. Requirements"))

intro_elements = []
intro = doc.add_heading("Introduction: Most significant new features", level=1)
intro_elements.append(intro._element)
lead = doc.add_paragraph(
    "This manual concentrates on the operational changes that most distinguish this AutoISF fork from stock AAPS and that are most likely to affect daily use:"
)
intro_elements.append(lead._element)
features = [
    "Self-healing coded Automation States.",
    "Settings import safeguards, including the four Keep current options.",
    "Tier 3 UAM Boost criteria, output limits and duration.",
    "BolusGiven strong-boost effects.",
    "BolusGivenMild and Mild Failsafe behavior.",
    "Steroid-management initiation and escalation.",
    "The Mounjaro three-day reduced-insulin cycle.",
    "Hypoglycaemia warnings and protective overrides.",
    "Bolus-calculator HP protection, split/delayed dosing, Lists 1 and 2, and the AIV history/export workflow."
]
intro_num_id = create_decimal_numbering(doc)
for feature in features:
    p = doc.add_paragraph(feature, style="Normal")
    apply_numbering(p, intro_num_id)
    intro_elements.append(p._element)
for element in intro_elements:
    first_section._element.addprevious(element)

doc.save(INTERMEDIATE)

# NAV_INPUT is produced by internal_nav.py. This second phase moves its generated
# Contents block from before the title to after the introduction and before section 1.
def move_contents_after_introduction():
    nav = Document(NAV_INPUT)
    title = next(p for p in nav.paragraphs if p.style.name == "Title")
    first = next(p for p in nav.paragraphs if p.style.name == "Heading 1" and p.text.startswith("1. Requirements"))
    toc_elements = []
    for child in list(nav._element.body):
        if child is title._element:
            break
        toc_elements.append(child)
    for child in toc_elements:
        first._element.addprevious(child)
    nav.save(OUTPUT)


if __name__ == "__main__":
    print(INTERMEDIATE)
