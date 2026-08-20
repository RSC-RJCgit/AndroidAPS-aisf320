from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 progress and hypo outcomes.docx")
doc = Document(path)
bookmarks = {n.get(qn("w:name")) for n in doc.element.body.xpath(".//w:bookmarkStart")}
anchors = [n.get(qn("w:anchor")) for n in doc.element.body.xpath(".//w:hyperlink[@w:anchor]")]
headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
unbookmarked = [p.text for p in headings if not p._p.xpath(".//w:bookmarkStart")]
required = [
    "Hypoglycaemia warning causes and outcomes.",
    "10. Factors affecting progress and hypoglycaemia outcomes",
    "a. Pod age and pump-type qualification",
    "b. Projected HP at bolus-calculator initiation",
    "c. Effects of 50recent, a 50% profile and a hypo temporary target on progress",
    "d. Hypoglycaemia warning causes and outcomes",
]
text = "\n".join(p.text for p in doc.paragraphs)
print({
    "tables": len(doc.tables),
    "headings": len(headings),
    "bookmarks": len(bookmarks),
    "hyperlinks": len(anchors),
    "broken_anchors": sorted(set(anchors) - bookmarks),
    "unbookmarked_headings": unbookmarked,
    "required_missing": [item for item in required if item not in text],
    "back_to_toc_count": text.count("Back to TOC"),
})
