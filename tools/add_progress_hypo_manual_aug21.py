from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\winword\aaa\AutoISF Idiosyncrasies and Operations Manual mydoc Aug 21 26 large linked introduction_Aug 21 26_DELL17.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF Operations Manual mydoc Aug 21 26 progress and hypo outcomes.docx")

doc = Document(SOURCE)


def next_bookmark_id():
    ids = []
    for node in doc.element.body.xpath(".//w:bookmarkStart"):
        try:
            ids.append(int(node.get(qn("w:id"))))
        except (TypeError, ValueError):
            pass
    return max(ids, default=0) + 1


bookmark_id = next_bookmark_id()


def add_bookmark(paragraph, name):
    global bookmark_id
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(1, start)
    paragraph._p.append(end)
    bookmark_id += 1


# The user's renamed end Contents heading lost its old bookmark during the earlier rename.
appendix_b_heading = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text == "Appendix B — Contents")
if not appendix_b_heading._p.xpath(".//w:bookmarkStart"):
    add_bookmark(appendix_b_heading, "appendix_b_contents")


def add_internal_link(paragraph, text, anchor, size_half_points=None):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    if size_half_points:
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(size_half_points))
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), str(size_half_points))
        props.extend([size, size_cs])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([props, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_paragraph_before(reference, text="", style=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    reference._element.addprevious(paragraph._element)
    return paragraph


appendix_a = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Appendix A"))

section = insert_paragraph_before(appendix_a, "10. Factors affecting progress and hypoglycaemia outcomes", "Heading 1")
add_bookmark(section, "sec10_progress_outcomes")

pod = insert_paragraph_before(appendix_a, "a. Pod age and pump-type qualification", "Heading 2")
add_bookmark(pod, "sec10_pod_age")
insert_paragraph_before(
    appendix_a,
    "Pod-age-dependent paths use the time since the last cannula/site-change record only after confirming that the active pump reports itself as a patch pump. On a tubed pump or MDI setup the pod-age helper returns no value, so a normal infusion-site age cannot accidentally satisfy a fresh-pod or old-pod condition. For a patch pump, age can affect progression in several distinct ways: a pod older than 60 hours can widen the NewDay2 calibration tier when the other sensor-adjustment gates are present; sustained high glucose for two hours after 60 hours can produce the one-per-pod old-pod warning; and fresh/old-pod automation branches can alter profile, target or weighting only when all their additional glucose, state, bolus-age and timing gates also pass. The pump-type check is therefore required and is already performed centrally before any of these calculations receives a pod age.",
    "Normal",
)

hp = insert_paragraph_before(appendix_a, "b. Projected HP at bolus-calculator initiation", "Heading 2")
add_bookmark(hp, "sec10_bolus_hp")
insert_paragraph_before(
    appendix_a,
    "The Wizard calculation already includes a projected-HP safety assessment during each live recalculation, including the initial display and Quick Wizard calculations that use the same BolusWizard engine. It is a prospective bolus-specific HP, not the live APS HP2 value: projected HP = current glucose - positive IOB - the percentage-adjusted proposed bolus, plus only negative current and short-average delta contributions at one quarter weight. When projected HP is at most 6.5 mmol/L, glucose is below 10.0 mmol/L, both rise measures are at least 0.6 mmol/L per five minutes, positive IOB is at least 2.0 U, and the proposal is at least 0.75 U, COB-derived insulin is removed and the remainder is reduced to 75%. The recalculated lower dose is shown directly as the proposed dose before confirmation; it is not merely a warning offering a later manual reduction.",
    "Normal",
)

recent = insert_paragraph_before(appendix_a, "c. Effects of 50recent, a 50% profile and a hypo temporary target on progress", "Heading 2")
add_bookmark(recent, "sec10_recent50_progress")
insert_paragraph_before(
    appendix_a,
    "These are related but not interchangeable states. An active 50% profile changes the effective insulin-to-carbohydrate calculation and can arm the automatic delayed-bolus review. LowBG=50recent independently halves only the newly entered carbohydrate component when the active profile is not already 50%, preventing the two mechanisms from stacking into a quarter dose; it also activates the recent-low rebound protection used by SMB calculations. Glucose below 5.0 mmol/L and not rising can invoke the same carbohydrate-only halving even without the state. A hypo temporary target changes the BG-correction target and can block, postpone or cancel later split/delayed delivery when the follow-up safety gates no longer pass. During delayed progress, the active 50% profile or 50recent state can preserve eligibility for reassessment, but every later dose is recalculated against fresh glucose, profile, IOB, pump availability, newer treatments and the relevant deadline; none of these states guarantees that the remainder will be delivered.",
    "Normal",
)

hypo = insert_paragraph_before(appendix_a, "d. Hypoglycaemia warning causes and outcomes", "Heading 2")
add_bookmark(hypo, "sec10_hypo_causes_outcomes")
insert_paragraph_before(
    appendix_a,
    "Warnings arise from several coded paths rather than one glucose threshold. GentleHypoRisk can be caused by agreement between low AAPS glucose and low/falling UKF raw glucose, by a falling 50% profile condition, or by HP2 at or below 4.1, with the additional requirement that the raw one- and five-minute directions are non-rising and that a low/HP2/activity qualifier is present. AlarmHypo1 covers a slow decline below 4.3 mmol/L, an emergency glucose below 3.0, a daytime steps-associated decline, or agreement between tightened HP2 and HP1 limits. AlarmHypo2 covers non-rising glucose at or below 4.3, glucose at or below 5.5 with substantial recent steps, or the same HP1/HP2 agreement. Outcomes are protective state and dosing-context changes as well as alerts: acceleration weighting is reduced, BGLstate is marked as recently low, AlarmHypo1/2 set LowBG=50recent, GentleHypoRisk lowers the IOB-threshold percentage, graph/CarePortal markers are written, and configured notifications/SMS messages are issued. Alert sound/SMS behavior can be suppressed in quiet hours while the protective state changes still occur, so the absence of an overnight audible warning does not mean that the hypo pathway did not act.",
    "Normal",
)

# Add the requested linked Introduction item while retaining the existing real numbering definition.
first_section = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("1. Requirements"))
first_section_index = next(i for i, p in enumerate(doc.paragraphs) if p._element is first_section._element)
intro_numbered = [p for p in doc.paragraphs[:first_section_index] if p._p.xpath("./w:pPr/w:numPr")]
if not intro_numbered:
    raise RuntimeError("Introduction numbered list not found")
intro_item = doc.add_paragraph(style="Normal")
intro_item._p.get_or_add_pPr().append(deepcopy(intro_numbered[0]._p.xpath("./w:pPr/w:numPr")[0]))
intro_item.paragraph_format.space_after = Pt(3)
add_internal_link(intro_item, "Hypoglycaemia warning causes and outcomes.", "sec10_hypo_causes_outcomes", 24)
first_section._element.addprevious(intro_item._element)

# Extend the linked end Contents immediately before its Appendix A entry.
contents_appendix_a = next(
    p for p in doc.paragraphs
    if p.text.startswith("Appendix A") and p._element is not appendix_a._element and p.style.name == "Normal"
)
contents_entries = [
    ("10. Factors affecting progress and hypoglycaemia outcomes", "sec10_progress_outcomes", False),
    ("a. Pod age and pump-type qualification", "sec10_pod_age", True),
    ("b. Projected HP at bolus-calculator initiation", "sec10_bolus_hp", True),
    ("c. Effects of 50recent, a 50% profile and a hypo temporary target on progress", "sec10_recent50_progress", True),
    ("d. Hypoglycaemia warning causes and outcomes", "sec10_hypo_causes_outcomes", True),
]
for text, anchor, indented in contents_entries:
    p = insert_paragraph_before(contents_appendix_a, style="Normal")
    if indented:
        p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    add_internal_link(p, text, anchor)

doc.save(OUTPUT)
print(OUTPUT)
