from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 carbohydrate graphs explained.docx")
OUTPUT = Path("AutoISF Operations Manual mydoc Aug 21 26 dura algorithm modifications explained.docx")
ANCHOR_TEXT = "d. TDD factor / sensitivity toggles"


document = Document(SOURCE)
anchor = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == ANCHOR_TEXT)

# Keep the alphabetic sequence correct after adding a new subsection d.
anchor.text = "e. TDD factor / sensitivity toggles"
overnight = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "e. The overnight safety-guard architecture")
overnight.text = "f. The overnight safety-guard architecture"


def insert_before(text: str, style: str = "Normal", bold_lead: str | None = None):
    paragraph = document.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    if style == "Normal":
        paragraph.paragraph_format.left_indent = Inches(0.5)
    else:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.keep_with_next = True
    anchor._p.addprevious(paragraph._p)
    return paragraph


insert_before("d. Duration (dura) modifications to the AutoISF algorithm", "Heading 2")
insert_before(
    "What dura measures. The glucose-status calculator walks backwards through valid CGM records and extends a plateau while each older reading remains within plus or minus 5% of the running average. It stops at the first reading outside that band, an unusable/gap-filled reading, or a gap greater than 13 minutes. The outputs are dura_ISF_minutes and the average glucose over that plateau; this is persistence near one level, not simply time above target and not the temporary-target duration.",
    bold_lead="What dura measures."
)
insert_before(
    "Engagement gates. Dura contributes no adaptation until the plateau has lasted at least 10 minutes, and it is bypassed whenever the plateau average is at or below the current target. Once both conditions are satisfied it is treated as prolonged high glucose that the existing profile ISF has not resolved.",
    bold_lead="Engagement gates."
)
insert_before(
    "Core calculation. Before the IOB safety taper, duraBoost = (dura minutes / 60) x (dura_ISF_weight / target BG) x (plateau-average BG - target BG), and dura_ISF = 1 + duraBoost. BG and target use the same mg/dL scale internally. A larger or longer elevation therefore raises the factor. In this algorithm a larger factor makes the effective ISF numerically smaller (profile ISF divided by the final factor), so dura is an insulin-strengthening/resistance response, not a sensitivity back-off.",
    bold_lead="Core calculation."
)
insert_before(
    "IOB-based safety taper. The current implementation applies the taper to duraBoost rather than reducing the configured weight globally. At IOB up to 1.5 U the full boost remains. Between 1.5 U and 3.0 U the multiplier falls linearly from 1.0 to 0.30. At or above 3.0 U the multiplier remains 0.30, so dura is trimmed but never completely disabled. If the current HP2 projection is above 7.0 mmol/L, the taper is deliberately bypassed and the full dura boost is retained. Script Debug records the latest taper time, IOB, multiplier and before/after boost, including on later cycles when no new taper occurs.",
    bold_lead="IOB-based safety taper."
)
insert_before(
    "Combination with the other AutoISF factors. When any strengthening factor is active, the algorithm initially takes the maximum of dura, BG, acceleration and postprandial factors. If acceleration is already braking below 1.0, that braking factor then weakens the selected maximum. The result passes through the ordinary AutoISF limits and sensitivity-ratio logic. If AutoISF weights are in display-only mode, dura is still calculated and graphed but is not applied to the returned dosing ISF.",
    bold_lead="Combination with the other AutoISF factors."
)
insert_before(
    "Settings and List 1 control. dura_ISF_weight is the live weight used by the calculation. Resting duration ISF weight (default 1.2, permitted range 0.0-3.0) stores the user's normal value. List 1 codes 5.022 and 5.024 change the resting value by -0.1 or +0.1. They also move the live value only when it still equals the former resting value; this prevents a separately changed live value from being overwritten. There is currently no distinct boosted/high dura-weight setting and no coded automation that automatically raises the live duration weight.",
    bold_lead="Settings and List 1 control."
)
insert_before(
    "Other decisions that consume dura. BolusGiven's manual-bolus branches require dura_ISF to be lower than acceleration ISF, so a duration-dominant rise does not enter those branches. The Virtual-Pump-only pseudo-wizard assessment measures how long dura and acceleration have remained active. OvernightDuraRescue can temporarily select the Standard profile for 60 minutes between 02:00 and 04:00 when the Low profile is active, glucose is above 6.0 mmol/L, dura and final ISF are both above 2.5, dura dominates every other factor and final ISF, no SMB was delivered in 30 minutes, no 50recent state exists, and both short- and long-average deltas are within approximately plus or minus 0.1 mmol/L. Its 60-minute cooldown prevents stacking.",
    bold_lead="Other decisions that consume dura."
)
insert_before(
    "Display and records. Each AIV result stores the calculated dura factor and final factor. The selectable DUR ISF graph shows its history, while Graph 5's pp/acce/du annotation row shows the current duration contribution together with the other principal weights. The target-offset/last-duration-taper annotation and Script Debug text help distinguish an actively tapered dura episode from a large untapered duration factor.",
    bold_lead="Display and records."
)

document.save(OUTPUT)
print(OUTPUT.resolve())
