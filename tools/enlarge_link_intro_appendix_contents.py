from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(r"C:\winword\aaa\AutoISF Idiosyncrasies and Operations Manual mydoc Aug 21 26 GUI and contents at end.docx")
OUTPUT = Path(r"C:\Users\arjay\StudioProjects\AaAPS3422a320\AutoISF_Operations_Manual_mydoc_Aug21_2026_large_linked_intro.docx")


doc = Document(SOURCE)


def heading_bookmark(prefix):
    paragraph = next(p for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.startswith(prefix))
    bookmark = paragraph._p.xpath(".//w:bookmarkStart")
    if not bookmark:
        raise RuntimeError(f"No bookmark for heading: {prefix}")
    return bookmark[0].get(qn("w:name"))


targets = {
    "bolus": heading_bookmark("a. Bolus calculator"),
    "lists": heading_bookmark("b. Lists 1 and 2"),
    "settings": heading_bookmark("2. Settings Import Features"),
    "aiv": heading_bookmark("c. AIV table"),
    "mechanism": heading_bookmark("6. Mechanism For Ignoring"),
    "self": heading_bookmark("b. Fixed 2026-08-19"),
    "tier3": heading_bookmark("b. Tier 3 UAM Boost"),
    "steroids": heading_bookmark("4. Setting: If Steroids"),
    "mj": heading_bookmark("5. Mounjaro"),
    "hypo": heading_bookmark("d. Hypoglycaemia warnings"),
    "gui": heading_bookmark("9. GUI display changes"),
}

intro_heading = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Introduction:"))
for run in intro_heading.runs:
    run.font.size = Pt(20)
    run.font.bold = True
intro_heading.paragraph_format.space_after = Pt(8)

lead = next(p for p in doc.paragraphs if p.text.startswith("This manual concentrates"))
for run in lead.runs:
    run.font.size = Pt(12)
lead.paragraph_format.space_after = Pt(6)

first_section = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("1. Requirements"))
old_intro_items = []
collecting = False
for p in doc.paragraphs:
    if p._element is lead._element:
        collecting = True
        continue
    if p._element is first_section._element:
        break
    if collecting and p._p.xpath("./w:pPr/w:numPr"):
        old_intro_items.append(p)

if not old_intro_items:
    raise RuntimeError("Introduction numbered list not found")
num_pr = deepcopy(old_intro_items[0]._p.xpath("./w:pPr/w:numPr")[0])
for p in old_intro_items:
    p._element.getparent().remove(p._element)

items = [
    ("Full description of the bolus calculator changes.", targets["bolus"]),
    ("Lists 1 and 2: full descriptions.", targets["lists"]),
    ("Settings Import Features — Especially With States.", targets["settings"]),
    ("AIV table.", targets["aiv"]),
    ("Mechanism For Ignoring Or Invoking Native Automations.", targets["mechanism"]),
    ("This is now self-healing for the coded automations' own states.", targets["self"]),
    ("Tier 3 UAM Boost criteria, output limits and duration.", targets["tier3"]),
    ("BolusGiven strong-boost effects.", targets["tier3"]),
    ("BolusGivenMild and Mild Failsafe behavior.", targets["tier3"]),
    ("Steroid-management initiation and escalation.", targets["steroids"]),
    ("The Mounjaro three-day reduced-insulin cycle.", targets["mj"]),
    ("Hypoglycaemia warnings and protective overrides.", targets["hypo"]),
    ("GUI display changes.", targets["gui"]),
]


def add_linked_numbered_item(text, anchor):
    p = doc.add_paragraph(style="Normal")
    p._p.get_or_add_pPr().append(deepcopy(num_pr))
    p.paragraph_format.space_after = Pt(3)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz"); size.set(qn("w:val"), "24")
    size_cs = OxmlElement("w:szCs"); size_cs.set(qn("w:val"), "24")
    r_pr.extend([color, underline, size, size_cs])
    text_node = OxmlElement("w:t"); text_node.text = text
    run.extend([r_pr, text_node]); hyperlink.append(run); p._p.append(hyperlink)
    first_section._element.addprevious(p._element)


for text, anchor in items:
    add_linked_numbered_item(text, anchor)

# The end-positioned navigation table is now explicitly an appendix.
contents = next(p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text == "Contents")
contents.text = "Appendix B — Contents"
contents.style = doc.styles["Heading 1"]

doc.save(OUTPUT)
print(OUTPUT)
