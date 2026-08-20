from pathlib import Path

from docx import Document
from docx.shared import Inches


SOURCE = Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 complete Settings Automations final.docx")
OUTPUT = Path("AutoISF Operations Manual mydoc Aug 21 26 List confirmation popups documented.docx")
AUTOMATION_HEADING = "AutoISF Coded Automations - Current Code Registry"


document = Document(SOURCE)
anchor = next(p for p in document.paragraphs if p.text.strip() == AUTOMATION_HEADING)


def insert_before(text: str, style: str = "Normal", bold_lead: str | None = None):
    paragraph = document.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    if style == "Normal":
        paragraph.paragraph_format.left_indent = Inches(0.25)
    anchor._p.addprevious(paragraph._p)
    return paragraph


heading = insert_before("c. Confirmation popups, current values and checkboxes", "Heading 2")
heading.paragraph_format.left_indent = Inches(0.25)

insert_before(
    "Opening the lists. Double-tap the IOB graph area for List 1. Double-tap the basal-rate icon area for List 2. "
    "On AAPSClient the list titles identify that commands will be relayed to the pump; on Pump/Virtual they identify direct local operation.",
    bold_lead="Opening the lists."
)
insert_before(
    "List 1 single commands. Selecting a toggle or one-shot row opens a separate confirmation popup. Where a live setting exists, the popup shows its current value before OK is pressed. "
    "For clean-graph and cloud-log commands it shows the effect instead, because there is no persistent value to display. OK applies the command; Cancel, the Android Back action, or tapping outside returns to List 1 without applying it.",
    bold_lead="List 1 single commands."
)
insert_before(
    "List 1 stepped commands. The confirmation popup shows the current value and two checkboxes labelled with the actual alternatives, such as -0.1/+0.1, -5%/+5%, Standard/Low, or NOMJremains/MJ3. "
    "The boxes are mutually exclusive: checking one clears the other. The user may correct the selection before pressing OK. OK applies only the checked choice; OK with neither box checked is a no-op. Cancel, Back, or outside dismissal returns to List 1.",
    bold_lead="List 1 stepped commands."
)
insert_before(
    "Which current value is displayed. Pump and Virtual builds read the local preference or state. AAPSClient displays the most recent settings snapshot received from the pump, identifies it as Pump current, and includes the age of that snapshot. "
    "If the pump has not supplied the relevant value, the popup reports Pump current: unavailable. The paired SMB-delivery row shows both baseline and mild-boost values; profile and MJ rows show their current named state rather than implying a numeric setting.",
    bold_lead="Which current value is displayed."
)
insert_before(
    "List 2 action commands. Selecting an MJ, steroid or AnyDesk action opens an OK/Cancel confirmation before execution; these action confirmations do not use setting checkboxes. Cancel returns to List 2. "
    "For AAPSClient AnyDesk, the confirmation also states whether an existing temporary target will be preserved and the command sent by NS Note only, or whether an NS Note plus five-minute relay TT will be used.",
    bold_lead="List 2 action commands."
)
insert_before(
    "List 2 graph controls. Selecting a graph row opens a popup with checkboxes already set to the current local display preferences. Graph on controls visibility. UKF1 and UKF3 also offer Use libre slope & offset. "
    "UKF2 shows that calibration box checked but disabled because calibration is already applied upstream. Graph5 instead uses the second box for BGL only (hide IA/carbs x3). OK saves the displayed choices and refreshes the Overview; Cancel, Back, or outside dismissal returns to List 2 without saving.",
    bold_lead="List 2 graph controls."
)
insert_before(
    "Confirmation and delivery status. Choosing and confirming a List 1 direction or List 2 action is the authorization step; the outer list itself does not execute a row. On Pump/Virtual, the corresponding handler is queued locally. "
    "On AAPSClient, dosing-setting/action commands use the documented relay route where applicable, while List 2 graph checkboxes remain local display settings and are never sent as relay temporary targets.",
    bold_lead="Confirmation and delivery status."
)

document.save(OUTPUT)
print(OUTPUT.resolve())
