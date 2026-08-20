from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\winword\aaa\AutoISF Operations Manual mydoc Aug 21 26 dura algorithm modifications explained.docx")
OUTPUT = Path("AutoISF Operations Manual mydoc Aug 21 26 additions moved from Contents to main.docx")

DURA_TITLE = "d. Duration (dura) modifications to the AutoISF algorithm"
CARB_TITLE = "d. What the different carbohydrate graphs mean"


document = Document(SOURCE)


def exact(text: str, style: str | None = None):
    matches = [p for p in document.paragraphs if p.text.strip() == text and (style is None or p.style.name == style)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one match for {text!r} style={style!r}; found {len(matches)}")
    return matches[0]


def toc_entry_before(paragraph, text: str):
    entry = document.add_paragraph(style="Normal")
    entry.add_run("  " + text)
    paragraph._p.addprevious(entry._p)


# The full additions were accidentally placed in the static opening Contents.
# Leave concise Contents entries in those locations, then move the existing OOXML
# paragraphs (preserving runs and formatting) into their genuine main sections.
dura_heading = exact(DURA_TITLE, "Heading 2")
dura_paragraphs = [dura_heading]
cursor = next(i for i, p in enumerate(document.paragraphs) if p._p is dura_heading._p) + 1
while cursor < len(document.paragraphs):
    paragraph = document.paragraphs[cursor]
    dura_paragraphs.append(paragraph)
    if paragraph.text.startswith("Display and records."):
        break
    cursor += 1
else:
    raise RuntimeError("Could not find the end of the misplaced dura section")

toc_entry_before(dura_heading, DURA_TITLE)
main_tdd = exact("d. TDD factor / sensitivity toggles", "Heading 2")
main_tdd.text = "e. TDD factor / sensitivity toggles"
main_overnight = exact("e. The overnight safety-guard architecture", "Heading 2")
main_overnight.text = "f. The overnight safety-guard architecture"
for paragraph in dura_paragraphs:
    main_tdd._p.addprevious(paragraph._p)

carb_heading = exact(CARB_TITLE, "Heading 2")
carb_paragraphs = [carb_heading]
cursor = next(i for i, p in enumerate(document.paragraphs) if p._p is carb_heading._p) + 1
while cursor < len(document.paragraphs):
    paragraph = document.paragraphs[cursor]
    carb_paragraphs.append(paragraph)
    if paragraph.text.startswith("Where the curves appear."):
        break
    cursor += 1
else:
    raise RuntimeError("Could not find the end of the misplaced carbohydrate section")

toc_entry_before(carb_heading, CARB_TITLE)
main_section_10 = exact("10. Factors affecting progress and hypoglycaemia outcomes", "Heading 1")
for paragraph in carb_paragraphs:
    main_section_10._p.addprevious(paragraph._p)

document.save(OUTPUT)
print(OUTPUT.resolve())
