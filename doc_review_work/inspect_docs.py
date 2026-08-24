from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent

for path in sorted(ROOT.glob("*.docx")):
    doc = Document(path)
    out = ROOT / f"{path.stem}.structure.txt"
    lines = [f"FILE: {path.name}", f"PARAGRAPHS: {len(doc.paragraphs)}", f"TABLES: {len(doc.tables)}", ""]
    for i, p in enumerate(doc.paragraphs):
        text = " ".join(p.text.split())
        if text:
            lines.append(f"P{i:04d}\t{p.style.name}\t{len(text):04d}\t{text}")
    for ti, table in enumerate(doc.tables):
        lines.append(f"\nTABLE {ti}: {len(table.rows)}x{len(table.columns)}")
        for ri, row in enumerate(table.rows):
            vals = [" ".join(c.text.split()) for c in row.cells]
            lines.append(f"T{ti}R{ri}\t" + " | ".join(vals))
    out.write_text("\n".join(lines), encoding="utf-8")
    print(out.name)
