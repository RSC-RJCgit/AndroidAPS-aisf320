from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent

FILES = {
    "AutoISF Operations Manual mydoc 2editAug 24 26 15 appendix D E F reorg.docx":
        "AutoISF Operations Manual revised 24 Aug 2026.docx",
    "AutoISF Operations Manual PLAIN LANGUAGE SUMMARY 2.docx":
        "AutoISF Operations Manual Plain Language Summary revised 24 Aug 2026.docx",
}

# Changelog paragraph opening -> topical destination heading opening.
MOVES = [
    ("Also fixed 2026-08-19, same day:", "c. A second, related gap"),
    ('"Walking soon" checkbox', "a. Bolus calculator:"),
    ("Quick Wizard parity", "a. Bolus calculator:"),
    ("As of 2026-08-22, rawDelta1Mgdl", "a. Where this sits"),
    ("One further wrinkle worth flagging", 'f. What actually reaches the AIV "FastRise" column'),
    ("Stacking-type criteria", "b. Tier 3 UAM Boost:"),
    ('Changed 2026-08-23: "Profile (manual override)"', "b. Lists 1 and 2:"),
    ("Reinstated the same day.", "b. Lists 1 and 2:"),
    ('Also fixed 2026-08-23 (second pass): "Cloud logs upload"', "b. Lists 1 and 2:"),
    ('New view-only row, added 2026-08-23: "Live Libre slope/offset', "b. Lists 1 and 2:"),
    ("Fixed 2026-08-23: OK now also returns to List 1", "b. Lists 1 and 2:"),
    ("Added 2026-08-23: three numeric stepped rows", "b. Lists 1 and 2:"),
    ("Also added the same day, Settings-only", "d. UKF1 for live AutoISF dosing"),
    ("Two UX gaps fixed 2026-08-23", "b. Lists 1 and 2:"),
    ("Re-pick entry point, added 2026-08-23", "d. Coded profile roles"),
    ("What changed. The six steroid-escalation", "e. Steroid escalation as a third coded role"),
    ("Two real incidents from live AIV/UserEntries data", "d. Coded profile roles"),
    ("One more hardcoded literal remains outside", "b. BatteryOver1pc:"),
    ("Two fixed BGL-attached labels removed", "c. Information shown on each graph"),
    ("Diagnostic logging added 2026-08-23", "d. Entry/exit gating and interaction with LowRaw24Cal"),
    ("GUI toggle added 2026-08-23", "c. Two independent enable switches"),
    ("Real cross-device discrepancy found the same day", "d. Entry/exit gating and interaction with LowRaw24Cal"),
    ('New List 1 row, added 2026-08-23: "Live Libre slope/offset', "d. Entry/exit gating and interaction with LowRaw24Cal"),
    ("BG floor added 2026-08-23", "c. LowBG"),
    ("Two further List 2 additions the same day", "b. Lists 1 and 2:"),
    ("showProfileNamesPopup() grew from 2 spinners", "e. Steroid escalation as a third coded role"),
]

DATE_SUMMARIES = {
    "2026-08-19": "Automation-state bootstrapping was made self-healing and given evidence-based neutral defaults. Full operational detail is now incorporated into Section 1.",
    "2026-08-22": "Bolus and Quick Wizard controls, Fast-Rise signal gating, and FastRise reporting were updated. The current behaviour and historical-export cautions are incorporated into Sections 8 and 15.",
    "2026-08-23": "Tier 3 safeguards, List controls, coded profile roles, UKF1 options, SensorAge diagnostics, graph labels, battery recovery, and LowBG recovery were updated. Details are incorporated into their corresponding topical sections and Appendix F.",
}

# These entries contain evidence, safety rationale, or historical-export interpretation
# that the topical text currently sends the reader to the changelog to find. Routine UI
# change narration is removed from the dense appendix but not duplicated in the sections,
# which already describe the resulting current behaviour.
MOVE_IN_FULL = {
    "Also fixed 2026-08-19, same day:",
    "As of 2026-08-22, rawDelta1Mgdl",
    "One further wrinkle worth flagging",
    "Stacking-type criteria",
    "Two real incidents from live AIV/UserEntries data",
    "Real cross-device discrepancy found the same day",
    "BG floor added 2026-08-23",
}


def heading_level(paragraph):
    m = re.fullmatch(r"Heading ([1-9])", paragraph.style.name or "")
    return int(m.group(1)) if m else None


def split_readably(text, max_chars=680):
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z\"'(])", text.strip())
    chunks, current = [], ""
    for sentence in sentences:
        candidate = sentence if not current else current + " " + sentence
        if current and len(candidate) > max_chars:
            chunks.append(current)
            current = sentence
        else:
            current = candidate
    if current:
        chunks.append(current)
    return chunks


def new_para_before(anchor, style_ppr, text):
    p = OxmlElement("w:p")
    if style_ppr is not None:
        p.append(deepcopy(style_ppr))
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    anchor._p.addprevious(p)


def append_to_section(doc, heading_prefix, source):
    target = next((p for p in doc.paragraphs if heading_level(p) and p.text.strip().startswith(heading_prefix)), None)
    if target is None:
        raise RuntimeError(f"Target heading not found: {heading_prefix}")
    level = heading_level(target)
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p._p is target._p)
    anchor = None
    for p in paragraphs[start + 1:]:
        pl = heading_level(p)
        if pl is not None and pl <= level:
            anchor = p
            break
    if anchor is None:
        raise RuntimeError(f"No following heading after: {heading_prefix}")
    ppr = source._p.pPr
    for chunk in split_readably(source.text):
        new_para_before(anchor, ppr, chunk)


def set_plain_text(paragraph, text):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    paragraph._p.append(r)


def process(src, dst):
    doc = Document(src)
    changelog = next(p for p in doc.paragraphs if heading_level(p) == 1 and p.text.strip().startswith("Appendix B — Changelog"))
    changelog_level = heading_level(changelog)
    in_changelog = False
    changelog_paras = []
    for p in doc.paragraphs:
        if p._p is changelog._p:
            in_changelog = True
            continue
        if in_changelog:
            pl = heading_level(p)
            if pl is not None and pl <= changelog_level:
                break
            changelog_paras.append(p)

    intro = next((p for p in changelog_paras if not heading_level(p)), None)
    if intro is None:
        raise RuntimeError("Changelog introduction not found")
    set_plain_text(intro, "Concise dated index of significant changes. Detailed explanations have been integrated into the relevant topical sections so each feature is described where it is used.")

    for opening, destination in MOVES:
        source = next((p for p in changelog_paras if p.text.strip().startswith(opening)), None)
        if source is None:
            raise RuntimeError(f"Changelog paragraph not found: {opening}")
        if opening in MOVE_IN_FULL:
            append_to_section(doc, destination, source)
        source._element.getparent().remove(source._element)

    # Replace the now-empty dated blocks with short navigation summaries.
    for date, summary in DATE_SUMMARIES.items():
        date_heading = next(p for p in doc.paragraphs if heading_level(p) == 3 and p.text.strip().startswith(date))
        date_index = next(i for i, p in enumerate(doc.paragraphs) if p._p is date_heading._p)
        following = doc.paragraphs[date_index + 1]
        new_para_before(following, intro._p.pPr, summary)

    doc.save(dst)
    print(dst.name)


for source_name, output_name in FILES.items():
    process(ROOT / source_name, ROOT / output_name)
