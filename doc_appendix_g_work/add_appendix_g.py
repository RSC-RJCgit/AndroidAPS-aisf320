from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).parent
SETTINGS_PATH = ROOT / "AutoISF_settings_Live_20260824_155639.txt"

FILES = {
    "Full Manual.docx": "AutoISF Operations Manual revised 24 Aug 2026 mydoc with Appendix G.docx",
    "Plain Language Manual.docx": "AutoISF Operations Manual Plain Language Summary revised 24 Aug 2026 mydoc with Appendix G.docx",
    "Brief Manual.docx": "AutoISF Operations BRIEF Manual Plain Language Summary revised 24 Aug 2026 mydoc with Appendix G.docx",
}


def parse_settings(path):
    rows = []
    for raw in path.read_text(encoding="utf-8-sig").splitlines():
        line = raw.strip()
        if not line:
            continue
        if " = " in line:
            key, value = line.split(" = ", 1)
        elif "=" in line:
            key, value = line.split("=", 1)
        else:
            key, value = line, ""
        rows.append((key.strip(), value.strip()))
    return rows


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_twips):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            set_cell_width(cell, width)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7B7B7")
        borders.append(element)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_cell(cell, *, bold=False, size=9):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_together = True
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(size)


def insert_toc_entry(doc):
    candidates = [
        p for p in doc.paragraphs
        if p.style.name == "Normal" and p.text.strip().startswith("Appendix F — Automation States In Use")
    ]
    if not candidates:
        return
    anchor = candidates[0]
    new_p = deepcopy(anchor._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Appendix G — Current Live Settings Snapshot"
    run.append(text)
    new_p.append(run)
    anchor._p.addnext(new_p)


def add_appendix(source, output, settings):
    doc = Document(source)
    insert_toc_entry(doc)

    # Full/plain manuals end with an "End" marker; keep it after the new appendix.
    end_text = None
    if doc.paragraphs and doc.paragraphs[-1].text.strip().startswith("End"):
        end_text = doc.paragraphs[-1].text
        remove_paragraph(doc.paragraphs[-1])

    # The Brief manual has a trailing empty Heading 1 placeholder. Reuse its position.
    if doc.paragraphs and doc.paragraphs[-1].style.name == "Heading 1" and not doc.paragraphs[-1].text.strip():
        remove_paragraph(doc.paragraphs[-1])

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Appendix G — Current Live Settings Snapshot [↑ TOC]")

    intro = doc.add_paragraph()
    intro.add_run("Snapshot source: ").bold = True
    intro.add_run("AutoISF_settings_Live_20260824_155639.txt")
    intro.add_run("\nCaptured: ").bold = True
    intro.add_run("24 August 2026, 3:56 PM — MAIN_PHONE_LOCAL (Live), full configuration, version 3.4.2.6+aisf321UK_654.")

    note = doc.add_paragraph()
    note.add_run("Important: ").bold = True
    note.add_run("These are point-in-time values captured from the Live device, not recommended defaults. Later app activity, profile changes, imports, or manual edits may change them.")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Normal Table"
    table.cell(0, 0).text = "Setting"
    table.cell(0, 1).text = "Live value"
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    shade_cell(table.cell(0, 0), "D9EAF7")
    shade_cell(table.cell(0, 1), "D9EAF7")
    format_cell(table.cell(0, 0), bold=True)
    format_cell(table.cell(0, 1), bold=True)

    for key, value in settings:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        prevent_row_split(table.rows[-1])
        format_cell(cells[0])
        format_cell(cells[1])

    # 7.0-inch usable width with the manuals' 0.75-inch margins.
    set_table_geometry(table, [6900, 3180])
    set_table_borders(table)

    if end_text:
        doc.add_paragraph(end_text)

    doc.save(output)
    print(output.name)


settings = parse_settings(SETTINGS_PATH)
if len(settings) != 75:
    raise RuntimeError(f"Expected 75 settings rows, found {len(settings)}")

for source_name, output_name in FILES.items():
    add_appendix(ROOT / source_name, ROOT / output_name, settings)
