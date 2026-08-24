from pathlib import Path
from docx import Document

for path in sorted(Path(__file__).parent.glob("*Appendix G.docx")):
    doc = Document(path)
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.startswith("Appendix G —")]
    assert len(headings) == 1, (path.name, len(headings))
    table = doc.tables[-1]
    assert len(table.rows) == 76, (path.name, len(table.rows))
    assert [c.text for c in table.rows[0].cells] == ["Setting", "Live value"]
    values = {row.cells[0].text: row.cells[1].text for row in table.rows[1:]}
    assert values["main_phone_snapshot_time"] == "8/24/26 03:56 PM"
    assert values["configuration_profile"] == "Current ProfileReal"
    assert values["split_bolus_interval"] == "7"
    print(f"PASS {path.name}: Appendix G heading=1, settings=75, tables={len(doc.tables)}")
