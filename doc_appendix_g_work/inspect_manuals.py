from pathlib import Path
from docx import Document

for path in sorted(Path(__file__).parent.glob("*.docx")):
    doc = Document(path)
    print(f"\n{path.name}: paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    print("HEADINGS:")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading"):
            print(f"  {i}: {p.style.name}: {p.text[:120]}")
    print("LAST PARAGRAPHS:")
    for i, p in list(enumerate(doc.paragraphs))[-15:]:
        print(f"  {i}: {p.style.name}: {p.text[:160]}")
